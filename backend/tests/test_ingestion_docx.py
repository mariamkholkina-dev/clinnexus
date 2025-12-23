"""
Тесты для DOCX ingestion (Шаг 4).
"""
from __future__ import annotations

import tempfile
from pathlib import Path
from uuid import UUID

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.enums import AnchorContentType, DocumentLifecycleStatus, DocumentType, IngestionStatus, StudyStatus
from app.db.models.anchors import Anchor
from app.db.models.auth import Workspace
from app.db.models.facts import Fact, FactEvidence
from app.db.models.studies import Document, DocumentVersion, Study
from app.services.ingestion import IngestionService


class TestDocxIngestion:
    """Тесты для DOCX ingestion."""
    
    @pytest.fixture
    async def test_workspace(self, db: AsyncSession) -> Workspace:
        """Создает тестовый workspace."""
        workspace = Workspace(name="Test Workspace")
        db.add(workspace)
        await db.commit()
        await db.refresh(workspace)
        return workspace
    
    @pytest.fixture
    async def test_study(self, db: AsyncSession, test_workspace: Workspace) -> Study:
        """Создает тестовое исследование."""
        study = Study(
            workspace_id=test_workspace.id,
            study_code="TEST-001",
            title="Test Study",
            status=StudyStatus.ACTIVE,
        )
        db.add(study)
        await db.commit()
        await db.refresh(study)
        return study
    
    @pytest.fixture
    async def test_document(self, db: AsyncSession, test_study: Study) -> Document:
        """Создает тестовый документ."""
        document = Document(
            workspace_id=test_study.workspace_id,
            study_id=test_study.id,
            doc_type=DocumentType.PROTOCOL,
            title="Test Document",
            lifecycle_status=DocumentLifecycleStatus.DRAFT,
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        return document
    
    @pytest.fixture
    def sample_docx_file(self) -> Path:
        """
        Создает тестовый DOCX файл с структурой:
        - Heading 1 "Introduction"
        - Paragraph "Protocol Version: 2.0"
        - Paragraph "Amendment Date: 05 March 2021"
        - Heading 2 "Objectives"
        - List item 1
        - List item 2
        """
        # Создаем временный файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        # Создаем документ через python-docx
        doc = DocxDocument()
        
        # Heading 1
        heading1 = doc.add_paragraph("Introduction", style="Heading 1")
        
        # Paragraph
        para1 = doc.add_paragraph("Protocol Version: 2.0")
        para2 = doc.add_paragraph("Amendment Date: 05 March 2021")
        
        # Heading 2
        heading2 = doc.add_paragraph("Objectives", style="Heading 2")
        
        # List items
        list_item1 = doc.add_paragraph("First objective", style="List Bullet")
        list_item2 = doc.add_paragraph("Second objective", style="List Bullet")
        
        # Сохраняем документ
        doc.save(str(tmp_path))
        
        return tmp_path
    
    @pytest.fixture
    async def test_version_with_docx(
        self, db: AsyncSession, test_document: Document, sample_docx_file: Path
    ) -> DocumentVersion:
        """Создает версию документа с DOCX файлом."""
        # Используем абсолютный путь как file:// URI
        # На Windows: file:///C:/path/to/file
        # На Unix: file:///path/to/file
        abs_path = sample_docx_file.resolve()
        if abs_path.as_posix().startswith("/"):
            # Unix-подобная система
            file_uri = f"file://{abs_path.as_posix()}"
        else:
            # Windows: C:/path/to/file -> file:///C:/path/to/file
            file_uri = f"file:///{abs_path.as_posix()}"
        
        version = DocumentVersion(
            document_id=test_document.id,
            version_label="v1.0",
            source_file_uri=file_uri,
            source_sha256="test_hash",
            ingestion_status=IngestionStatus.UPLOADED,
        )
        db.add(version)
        await db.commit()
        await db.refresh(version)
        return version
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_creates_anchors(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что DOCX ingestion создаёт anchors."""
        service = IngestionService(db)
        result = await service.ingest(test_version_with_docx.id)
        
        # Проверяем результат
        assert result.anchors_created > 0
        assert result.doc_version_id == test_version_with_docx.id
        
        # Проверяем, что anchors созданы в БД
        stmt = select(Anchor).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchors_result = await db.execute(stmt)
        anchors = anchors_result.scalars().all()
        
        assert len(anchors) == result.anchors_created
        assert len(anchors) > 0
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_content_types(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что создаются anchors разных типов (hdr, p, li)."""
        service = IngestionService(db)
        await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        # Проверяем наличие разных типов
        stmt = select(Anchor).where(
            Anchor.doc_version_id == test_version_with_docx.id
        )
        anchors_result = await db.execute(stmt)
        anchors_by_type = {}
        for anchor in anchors_result.scalars().all():
            content_type = anchor.content_type
            anchors_by_type[content_type] = anchors_by_type.get(content_type, 0) + 1
        
        # Должны быть заголовки (hdr)
        assert AnchorContentType.HDR in anchors_by_type
        assert anchors_by_type[AnchorContentType.HDR] >= 2  # Introduction и Objectives
        
        # Должен быть параграф (p)
        assert AnchorContentType.P in anchors_by_type
        assert anchors_by_type[AnchorContentType.P] >= 2  # Protocol Version + Amendment Date
        
        # Должны быть элементы списка (li)
        assert AnchorContentType.LI in anchors_by_type
        assert anchors_by_type[AnchorContentType.LI] >= 2  # Два элемента списка
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_section_path(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что section_path корректно отражает структуру документа."""
        service = IngestionService(db)
        await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        # Получаем все anchors
        stmt = select(Anchor).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchors_result = await db.execute(stmt)
        anchors = anchors_result.scalars().all()
        
        # Собираем уникальные section_path
        section_paths = set(anchor.section_path for anchor in anchors)
        
        # Должны быть:
        # - "Introduction" (или "ROOT" если нет заголовков выше)
        # - "Introduction/Objectives" (заголовок второго уровня)
        
        # Проверяем наличие "Introduction" в путях
        intro_paths = [p for p in section_paths if "Introduction" in p]
        assert len(intro_paths) > 0, f"Не найдено 'Introduction' в путях: {section_paths}"
        
        # Проверяем наличие "Introduction/Objectives" (если заголовки правильно обработаны)
        objectives_paths = [p for p in section_paths if "Objectives" in p]
        assert len(objectives_paths) > 0, f"Не найдено 'Objectives' в путях: {section_paths}"
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_anchor_id_format(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест формата anchor_id."""
        service = IngestionService(db)
        await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        # Получаем anchors
        stmt = select(Anchor).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchors_result = await db.execute(stmt)
        anchors = anchors_result.scalars().all()
        
        doc_version_id_str = str(test_version_with_docx.id)
        
        for anchor in anchors:
            parts = anchor.anchor_id.split(":")
            
            if anchor.content_type == AnchorContentType.FN:
                # Формат для footnotes: {doc_version_id}:fn:{fn_index}:{fn_para_index}:{text_hash}
                assert len(parts) == 5, f"Неверный формат anchor_id для footnote: {anchor.anchor_id}"
                anchor_doc_version_id, fn_prefix, fn_index_str, fn_para_index_str, text_hash = parts
                
                # Проверяем префикс для footnotes
                assert fn_prefix == "fn", f"Ожидается 'fn' для footnote, получено: {fn_prefix}"
                
                # Проверяем индексы
                assert fn_index_str.isdigit()
                assert fn_para_index_str.isdigit()
                
                # Проверяем соответствие location_json
                assert anchor.location_json.get("fn_index") == int(fn_index_str)
                assert anchor.location_json.get("fn_para_index") == int(fn_para_index_str)
            else:
                # Формат для paragraph-anchors: {doc_version_id}:{content_type}:{para_index}:{text_hash}
                assert len(parts) == 4, f"Неверный формат anchor_id для paragraph: {anchor.anchor_id}"
                anchor_doc_version_id, content_type_str, para_index_str, text_hash = parts
                
                # Проверяем content_type
                assert content_type_str in ["p", "li", "hdr"]
                assert anchor.content_type.value == content_type_str
                
                # Проверяем para_index (должен быть числом)
                assert para_index_str.isdigit()
                para_index = int(para_index_str)
                assert para_index == anchor.location_json.get("para_index")
            
            # Проверяем doc_version_id
            assert anchor_doc_version_id == doc_version_id_str
            
            # Проверяем hash (64 символа hex)
            assert len(text_hash) == 64
            assert all(c in "0123456789abcdef" for c in text_hash)
            
            # Проверяем, что section_path НЕ входит в anchor_id
            assert anchor.section_path not in anchor.anchor_id, \
                f"section_path не должен быть частью anchor_id, но найден: {anchor.section_path}"
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_summary_counts(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что ingestion_summary_json содержит counts."""
        # Запускаем через endpoint для получения полного summary
        from app.api.v1.documents import start_ingestion
        
        result = await start_ingestion(test_version_with_docx.id, force=False, db=db)
        
        # Обновляем версию из БД
        await db.refresh(test_version_with_docx)
        
        summary = test_version_with_docx.ingestion_summary_json
        assert summary is not None
        
        # Проверяем наличие counts_by_type
        assert "counts_by_type" in summary
        counts_by_type = summary["counts_by_type"]
        assert isinstance(counts_by_type, dict)
        assert "hdr" in counts_by_type
        assert "p" in counts_by_type
        assert "li" in counts_by_type
        
        # Проверяем num_sections
        assert "num_sections" in summary
        assert summary["num_sections"] > 0
        
        # Проверяем sections
        assert "sections" in summary
        assert isinstance(summary["sections"], list)
        assert len(summary["sections"]) == summary["num_sections"]

        # Проверяем, что facts extraction тоже записан в summary
        assert "facts_extraction" in summary
        assert "facts_count" in summary["facts_extraction"]
        assert "needs_review" in summary["facts_extraction"]

        # Проверяем, что после ingest в БД есть protocol_meta факты и evidence с реальными anchor_id
        facts_stmt = select(Fact).where(Fact.created_from_doc_version_id == test_version_with_docx.id)
        facts = (await db.execute(facts_stmt)).scalars().all()
        assert any((f.fact_type, f.fact_key) == ("protocol_meta", "protocol_version") for f in facts)

        # evidence (если есть) не должен содержать "anchor_1" и должен ссылаться на anchors текущей версии
        anchors_stmt = select(Anchor.anchor_id).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchor_ids = {row[0] for row in (await db.execute(anchors_stmt)).all()}

        if facts:
            fact_ids = [f.id for f in facts]
            ev_stmt = select(FactEvidence).where(FactEvidence.fact_id.in_(fact_ids))
            evidence = (await db.execute(ev_stmt)).scalars().all()
            assert all(e.anchor_id != "anchor_1" for e in evidence)
            assert all(e.anchor_id in anchor_ids for e in evidence)
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_re_ingest(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что re-ingest пересоздаёт anchors без ошибок."""
        service = IngestionService(db)
        
        # Первый ingestion
        result1 = await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        anchors_count_1 = result1.anchors_created
        assert anchors_count_1 > 0
        
        # Второй ingestion (re-ingest)
        result2 = await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        # Проверяем, что количество anchors не изменилось (пересозданы)
        assert result2.anchors_created == anchors_count_1
        
        # Проверяем, что в БД правильное количество
        stmt = select(Anchor).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchors_result = await db.execute(stmt)
        anchors = anchors_result.scalars().all()
        assert len(anchors) == anchors_count_1
    
    @pytest.mark.asyncio
    async def test_docx_ingestion_location_json(
        self, db: AsyncSession, test_version_with_docx: DocumentVersion
    ):
        """Тест, что location_json содержит необходимые поля."""
        service = IngestionService(db)
        await service.ingest(test_version_with_docx.id)
        await db.commit()
        
        # Получаем anchors
        stmt = select(Anchor).where(Anchor.doc_version_id == test_version_with_docx.id)
        anchors_result = await db.execute(stmt)
        anchors = anchors_result.scalars().all()
        
        for anchor in anchors:
            location = anchor.location_json
            assert isinstance(location, dict)
            assert "para_index" in location
            assert "style" in location
            assert "section_path" in location
            assert isinstance(location["para_index"], int)
            assert location["para_index"] > 0
            assert location["section_path"] == anchor.section_path
    
    @pytest.mark.asyncio
    async def test_anchor_id_stability_with_section_path_change(
        self, db: AsyncSession, test_document: Document
    ):
        """
        Тест стабильности anchor_id при изменении section_path.
        
        Проверяет, что при одинаковом тексте и para_index anchor_id совпадает,
        даже если section_path отличается.
        """
        from app.services.ingestion.docx_ingestor import DocxIngestor
        from uuid import uuid4
        import tempfile
        
        # Создаём два документа с одинаковым содержимым параграфов,
        # но разной структурой заголовков (разный section_path)
        
        # Документ 1: "Section A" -> "Paragraph 1"
        tmp_file1 = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path1 = Path(tmp_file1.name)
        tmp_file1.close()
        
        doc1 = DocxDocument()
        doc1.add_paragraph("Section A", style="Heading 1")
        doc1.add_paragraph("Paragraph 1")  # para_index = 2
        doc1.save(str(tmp_path1))
        
        # Документ 2: "Section B" -> "Subsection" -> "Paragraph 1"
        # Тот же текст параграфа, но другой section_path
        tmp_file2 = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path2 = Path(tmp_file2.name)
        tmp_file2.close()
        
        doc2 = DocxDocument()
        doc2.add_paragraph("Section B", style="Heading 1")
        doc2.add_paragraph("Subsection", style="Heading 2")
        doc2.add_paragraph("Paragraph 1")  # para_index = 3 (но текст тот же)
        doc2.save(str(tmp_path2))
        
        try:
            doc_version_id = uuid4()
            ingestor = DocxIngestor()
            
            # Обрабатываем первый документ
            result1 = ingestor.ingest(tmp_path1, doc_version_id)
            
            # Обрабатываем второй документ (с другим doc_version_id, чтобы проверить логику)
            # Но на самом деле, нам нужно проверить, что при одинаковом para_index и тексте
            # anchor_id будет одинаковым. Для этого используем тот же doc_version_id,
            # но создадим документ с теми же параграфами, но разными заголовками.
            
            # Найдём параграф "Paragraph 1" в первом результате
            para1_anchor1 = None
            for anchor in result1.anchors:
                if anchor.text_norm == "Paragraph 1" and anchor.content_type == AnchorContentType.P:
                    para1_anchor1 = anchor
                    break
            
            assert para1_anchor1 is not None, "Не найден параграф 'Paragraph 1' в первом документе"
            
            # Проверяем, что anchor_id не содержит section_path
            assert para1_anchor1.section_path not in para1_anchor1.anchor_id
            
            # Проверяем формат: {doc_version_id}:{content_type}:{para_index}:{text_hash}
            parts = para1_anchor1.anchor_id.split(":")
            assert len(parts) == 4
            assert parts[0] == str(doc_version_id)
            assert parts[1] == AnchorContentType.P.value
            assert parts[2] == str(para1_anchor1.location_json["para_index"])
            
            # Теперь создадим документ, где параграф "Paragraph 1" имеет тот же para_index
            # (нужно создать документ без лишних параграфов перед ним)
            tmp_file3 = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp_path3 = Path(tmp_file3.name)
            tmp_file3.close()
            
            doc3 = DocxDocument()
            doc3.add_paragraph("Different Section", style="Heading 1")  # Другой заголовок
            doc3.add_paragraph("Paragraph 1")  # Тот же текст, тот же para_index = 2
            doc3.save(str(tmp_path3))
            
            # Используем другой doc_version_id, но с тем же текстом и para_index
            # anchor_id должен иметь тот же text_hash, но другой doc_version_id
            # Для полной проверки стабильности используем тот же doc_version_id
            result3 = ingestor.ingest(tmp_path3, doc_version_id)
            
            # Найдём параграф "Paragraph 1" во втором результате
            para1_anchor3 = None
            for anchor in result3.anchors:
                if anchor.text_norm == "Paragraph 1" and anchor.content_type == AnchorContentType.P:
                    para1_anchor3 = anchor
                    break
            
            assert para1_anchor3 is not None, "Не найден параграф 'Paragraph 1' во втором документе"
            
            # Проверяем, что para_index совпадает
            assert para1_anchor1.location_json["para_index"] == para1_anchor3.location_json["para_index"]
            
            # Проверяем, что text_hash совпадает (текст одинаковый)
            assert para1_anchor1.text_hash == para1_anchor3.text_hash
            
            # Проверяем, что anchor_id совпадает (при одинаковом doc_version_id, content_type, para_index и text_hash)
            assert para1_anchor1.anchor_id == para1_anchor3.anchor_id, \
                f"anchor_id должен совпадать при одинаковом тексте и para_index, " \
                f"но отличается: {para1_anchor1.anchor_id} vs {para1_anchor3.anchor_id}"
            
            # Проверяем, что section_path отличается (что и ожидается)
            assert para1_anchor1.section_path != para1_anchor3.section_path, \
                "section_path должен отличаться, но совпадает"
            
        finally:
            # Удаляем временные файлы
            if tmp_path1.exists():
                tmp_path1.unlink()
            if tmp_path2.exists():
                tmp_path2.unlink()
            if tmp_path3.exists():
                tmp_path3.unlink()

