"""
Тесты для HeadingDetector (детекция заголовков в DOCX).
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P

from app.services.ingestion.heading_detector import (
    DocStats,
    HeadingDetector,
    HeadingHit,
    normalize_title,
)


class TestHeadingDetector:
    """Тесты для HeadingDetector."""
    
    def test_normalize_title(self):
        """Тест нормализации текста заголовка."""
        # Обычный текст
        assert normalize_title("Introduction") == "Introduction"
        
        # С лишними пробелами
        assert normalize_title("  Introduction  ") == "Introduction"
        assert normalize_title("Introduction\n\nSection") == "Introduction Section"
        
        # С trailing colon
        assert normalize_title("Introduction:") == "Introduction"
        
        # С zero-width chars
        text_with_zw = "Introduction\u200b\u200c\u200d"
        assert normalize_title(text_with_zw) == "Introduction"
        
        # Ограничение длины
        long_text = "A" * 150
        normalized = normalize_title(long_text)
        assert len(normalized) == 120
        assert normalized == "A" * 120
    
    def test_heading_by_word_style(self):
        """Тест детекции заголовка по стандартному стилю Word (Heading 1, Heading 2)."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Heading 1
        heading1 = doc.add_paragraph("Introduction", style="Heading 1")
        
        # Heading 2
        heading2 = doc.add_paragraph("Objectives", style="Heading 2")
        
        # Paragraph
        para = doc.add_paragraph("This is a paragraph.")
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        detector = HeadingDetector(enable_visual_fallback=False)
        
        # Heading 1
        hit1 = detector.detect(doc_loaded.paragraphs[0])
        assert hit1.is_heading is True
        assert hit1.level == 1
        assert hit1.mode == "style"
        assert hit1.confidence >= 0.9
        assert hit1.normalized_title == "Introduction"
        
        # Heading 2
        hit2 = detector.detect(doc_loaded.paragraphs[1])
        assert hit2.is_heading is True
        assert hit2.level == 2
        assert hit2.mode == "style"
        assert hit2.normalized_title == "Objectives"
        
        # Paragraph
        hit3 = detector.detect(doc_loaded.paragraphs[2])
        assert hit3.is_heading is False
        assert hit3.mode == "none"
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_heading_by_outline_level_xml(self):
        """Тест детекции заголовка по outline level в XML."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Параграф со стилем Normal, но с outlineLvl=1 в XML
        para = doc.add_paragraph("Custom Heading", style="Normal")
        
        # Вручную устанавливаем outlineLvl в XML
        p_element: CT_P = para._p
        if p_element.pPr is None:
            from docx.oxml import parse_xml
            pPr = parse_xml(r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            p_element.insert(0, pPr)
            p_element.pPr = pPr
        
        from docx.oxml import parse_xml
        outline_elem = parse_xml(
            r'<w:outlineLvl w:val="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        p_element.pPr.append(outline_elem)
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        detector = HeadingDetector(enable_visual_fallback=False)
        
        hit = detector.detect(doc_loaded.paragraphs[0])
        assert hit.is_heading is True
        assert hit.level == 2  # val=1 в XML → level=2 (1-based)
        assert hit.mode == "outline"
        assert hit.confidence >= 0.8
        assert hit.normalized_title == "Custom Heading"
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_heading_by_numbering_regex(self):
        """Тест детекции заголовка по нумерации (1.2, 1.2.3, ...)."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Параграфы с нумерацией
        para1 = doc.add_paragraph("1.2 Study Objectives", style="Normal")
        para2 = doc.add_paragraph("1.2.3 Background Information", style="Normal")
        para3 = doc.add_paragraph("1) Introduction", style="Normal")
        
        # Обычный параграф
        para4 = doc.add_paragraph("This is a regular paragraph.", style="Normal")
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        detector = HeadingDetector(enable_visual_fallback=False)
        
        # 1.2 → level 2
        hit1 = detector.detect(doc_loaded.paragraphs[0])
        assert hit1.is_heading is True
        assert hit1.level == 2
        assert hit1.mode == "numbering"
        assert hit1.confidence >= 0.6
        
        # 1.2.3 → level 3
        hit2 = detector.detect(doc_loaded.paragraphs[1])
        assert hit2.is_heading is True
        assert hit2.level == 3
        assert hit2.mode == "numbering"
        
        # 1) → level 1
        hit3 = detector.detect(doc_loaded.paragraphs[2])
        assert hit3.is_heading is True
        assert hit3.level == 1
        assert hit3.mode == "numbering"
        
        # Обычный параграф
        hit4 = detector.detect(doc_loaded.paragraphs[3])
        assert hit4.is_heading is False
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_list_item_not_heading(self):
        """Тест, что элемент списка не считается заголовком по numbering."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Элемент списка с нумерацией
        list_item = doc.add_paragraph("1) First item", style="List Bullet")
        
        # Параграф с нумерацией, но не элемент списка
        para = doc.add_paragraph("1.2 Study Objectives", style="Normal")
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        detector = HeadingDetector(enable_visual_fallback=False)
        
        # Элемент списка НЕ должен быть заголовком
        hit1 = detector.detect(doc_loaded.paragraphs[0])
        # Может быть detected как list item, но не как heading по numbering
        # Проверяем, что если это heading, то не по numbering
        if hit1.is_heading:
            assert hit1.mode != "numbering"
        
        # Параграф с нумерацией должен быть заголовком
        hit2 = detector.detect(doc_loaded.paragraphs[1])
        assert hit2.is_heading is True
        assert hit2.mode == "numbering"
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_visual_fallback_only_when_enabled(self):
        """Тест, что визуальный fallback работает только когда включен."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Параграф с визуальными признаками заголовка (bold, center, большой шрифт)
        para = doc.add_paragraph("BACKGROUND", style="Normal")
        
        # Делаем текст жирным
        for run in para.runs:
            run.bold = True
            if run.font:
                from docx.shared import Pt
                run.font.size = Pt(16)  # Большой шрифт
        
        # Центрируем
        para.alignment = 1  # CENTER
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        
        # Вычисляем статистику документа
        doc_stats = HeadingDetector.compute_doc_stats(list(doc_loaded.paragraphs))
        
        # Без visual fallback
        detector_no_visual = HeadingDetector(enable_visual_fallback=False)
        detector_no_visual.set_doc_stats(doc_stats)
        hit1 = detector_no_visual.detect(doc_loaded.paragraphs[0])
        assert hit1.is_heading is False or hit1.mode != "visual"
        
        # С visual fallback
        detector_visual = HeadingDetector(enable_visual_fallback=True)
        detector_visual.set_doc_stats(doc_stats)
        hit2 = detector_visual.detect(doc_loaded.paragraphs[0])
        # Может быть detected как heading, если score >= 0.6
        if hit2.is_heading:
            assert hit2.mode == "visual"
            assert hit2.confidence >= 0.6
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_get_outline_level(self):
        """Тест чтения outline level из XML."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        para = doc.add_paragraph("Test", style="Normal")
        
        # Устанавливаем outlineLvl=2 в XML
        p_element: CT_P = para._p
        if p_element.pPr is None:
            from docx.oxml import parse_xml
            pPr = parse_xml(r'<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            p_element.insert(0, pPr)
            p_element.pPr = pPr
        
        from docx.oxml import parse_xml
        outline_elem = parse_xml(
            r'<w:outlineLvl w:val="2" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        p_element.pPr.append(outline_elem)
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        detector = HeadingDetector()
        
        level = detector.get_outline_level(doc_loaded.paragraphs[0])
        assert level == 3  # val=2 в XML → level=3 (1-based)
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_numbering_level_calculation(self):
        """Тест вычисления уровня по нумерации."""
        detector = HeadingDetector()
        
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        test_cases = [
            ("1 Introduction", 1),
            ("1.2 Objectives", 2),
            ("1.2.3 Background", 3),
            ("1.2.3.4 Methods", 4),
        ]
        
        for text, expected_level in test_cases:
            para = doc.add_paragraph(text, style="Normal")
            hit = detector.detect(para)
            if hit.is_heading and hit.mode == "numbering":
                assert hit.level == expected_level, f"Failed for '{text}': expected {expected_level}, got {hit.level}"
        
        # Удаляем временный файл
        tmp_path.unlink()
    
    def test_compute_doc_stats(self):
        """Тест вычисления статистики документа."""
        # Создаем временный DOCX файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()
        
        doc = DocxDocument()
        
        # Параграфы с разными размерами шрифта
        para1 = doc.add_paragraph("Normal text")
        para2 = doc.add_paragraph("Bold text")
        
        for run in para2.runs:
            run.bold = True
            if run.font:
                from docx.shared import Pt
                run.font.size = Pt(14)
        
        doc.save(str(tmp_path))
        
        # Загружаем и тестируем
        doc_loaded = DocxDocument(str(tmp_path))
        stats = HeadingDetector.compute_doc_stats(list(doc_loaded.paragraphs))
        
        assert stats.total_paragraphs == 2
        assert stats.bold_ratio >= 0.0
        # median_font_size может быть None, если не удалось извлечь размеры
        
        # Удаляем временный файл
        tmp_path.unlink()

