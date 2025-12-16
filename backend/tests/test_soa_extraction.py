"""
Тесты для SoA extraction (Шаг 5).
"""
from __future__ import annotations

import tempfile
from pathlib import Path
from uuid import UUID

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.enums import (
    AnchorContentType,
    DocumentLifecycleStatus,
    DocumentType,
    EvidenceRole,
    FactStatus,
    IngestionStatus,
    StudyStatus,
)
from app.db.models.anchors import Anchor
from app.db.models.auth import Workspace
from app.db.models.facts import Fact, FactEvidence
from app.db.models.studies import Document, DocumentVersion, Study
from app.services.ingestion import IngestionService


class TestSoAExtraction:
    """Тесты для SoA extraction."""

    @pytest.fixture
    async def test_workspace(self, db: AsyncSession) -> Workspace:
        """Создает тестовый workspace."""
        workspace = Workspace(name="Test Workspace")
        db.add(workspace)
        await db.commit()
        await db.refresh(workspace)
        return workspace

    @pytest.fixture
    async def test_study(self, db: AsyncSession, test_workspace: Workspace) -> Study:
        """Создает тестовое исследование."""
        study = Study(
            workspace_id=test_workspace.id,
            study_code="TEST-001",
            title="Test Study",
            status=StudyStatus.ACTIVE,
        )
        db.add(study)
        await db.commit()
        await db.refresh(study)
        return study

    @pytest.fixture
    async def test_document(self, db: AsyncSession, test_study: Study) -> Document:
        """Создает тестовый документ."""
        document = Document(
            workspace_id=test_study.workspace_id,
            study_id=test_study.id,
            doc_type=DocumentType.PROTOCOL,
            title="Test Protocol",
            lifecycle_status=DocumentLifecycleStatus.DRAFT,
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        return document

    @pytest.fixture
    def soa_docx_file(self) -> Path:
        """
        Создает тестовый DOCX файл с SoA таблицей:
        - Heading 1 "Schedule of Activities"
        - Таблица 4x4:
          [0,0]="Procedure", [0,1]="Screening", [0,2]="Baseline", [0,3]="Week 4"
          rows procedures: "Informed consent", "Vitals", "ECG"
          cells mark: X в нескольких местах
        """
        # Создаем временный файл
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()

        # Создаем документ через python-docx
        doc = DocxDocument()

        # Heading 1
        heading1 = doc.add_paragraph("Schedule of Activities", style="Heading 1")

        # Создаем таблицу 4x4
        table = doc.add_table(rows=4, cols=4)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Procedure"
        header_cells[1].text = "Screening"
        header_cells[2].text = "Baseline"
        header_cells[3].text = "Week 4"

        # Procedure rows
        proc_rows = [
            ("Informed consent", "X", "X", ""),
            ("Vitals", "X", "X", "X"),
            ("ECG", "", "X", ""),
        ]

        for i, (proc_name, *marks) in enumerate(proc_rows, start=1):
            row = table.rows[i]
            row.cells[0].text = proc_name
            for j, mark in enumerate(marks, start=1):
                row.cells[j].text = mark

        # Сохраняем документ
        doc.save(str(tmp_path))

        return tmp_path

    @pytest.fixture
    async def test_version_with_soa(
        self, db: AsyncSession, test_document: Document, soa_docx_file: Path
    ) -> DocumentVersion:
        """Создает версию документа с DOCX файлом, содержащим SoA."""
        # Используем абсолютный путь как file:// URI
        abs_path = soa_docx_file.resolve()
        if abs_path.as_posix().startswith("/"):
            # Unix-подобная система
            file_uri = f"file://{abs_path.as_posix()}"
        else:
            # Windows: C:/path/to/file -> file:///C:/path/to/file
            file_uri = f"file:///{abs_path.as_posix()}"

        version = DocumentVersion(
            document_id=test_document.id,
            version_label="v1.0",
            source_file_uri=file_uri,
            source_sha256="test_hash",
            ingestion_status=IngestionStatus.UPLOADED,
        )
        db.add(version)
        await db.commit()
        await db.refresh(version)
        return version

    @pytest.mark.asyncio
    async def test_soa_extraction_creates_cell_anchors(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест, что SoA extraction создаёт cell anchors."""
        service = IngestionService(db)
        result = await service.ingest(test_version_with_soa.id)
        await db.commit()

        # Проверяем результат
        assert result.soa_detected is True
        assert result.cell_anchors_created > 0

        # Проверяем, что cell anchors созданы в БД
        stmt = select(Anchor).where(
            Anchor.doc_version_id == test_version_with_soa.id,
            Anchor.content_type == AnchorContentType.CELL,
        )
        anchors_result = await db.execute(stmt)
        cell_anchors = anchors_result.scalars().all()

        assert len(cell_anchors) == result.cell_anchors_created
        assert len(cell_anchors) > 0

        # Проверяем формат anchor_id для cell anchors
        for anchor in cell_anchors:
            parts = anchor.anchor_id.split(":")
            assert len(parts) == 5
            assert parts[2] == "cell"  # content_type
            assert anchor.content_type == AnchorContentType.CELL

    @pytest.mark.asyncio
    async def test_soa_extraction_location_json(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест, что location_json для cell anchors содержит необходимые поля."""
        service = IngestionService(db)
        await service.ingest(test_version_with_soa.id)
        await db.commit()

        # Получаем cell anchors
        stmt = select(Anchor).where(
            Anchor.doc_version_id == test_version_with_soa.id,
            Anchor.content_type == AnchorContentType.CELL,
        )
        anchors_result = await db.execute(stmt)
        cell_anchors = anchors_result.scalars().all()

        assert len(cell_anchors) > 0

        for anchor in cell_anchors:
            location = anchor.location_json
            assert isinstance(location, dict)
            assert "table_index" in location
            assert "row_idx" in location
            assert "col_idx" in location
            assert "is_header" in location
            assert "header_path" in location
            assert isinstance(location["table_index"], int)
            assert isinstance(location["row_idx"], int)
            assert isinstance(location["col_idx"], int)
            assert isinstance(location["is_header"], bool)
            assert isinstance(location["header_path"], dict)
            assert "row_headers" in location["header_path"]
            assert "col_headers" in location["header_path"]

    @pytest.mark.asyncio
    async def test_soa_extraction_creates_facts(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест, что SoA extraction создаёт facts."""
        service = IngestionService(db)
        result = await service.ingest(test_version_with_soa.id)
        await db.commit()

        # Получаем document для study_id
        document = await db.get(Document, test_version_with_soa.document_id)
        assert document is not None

        # Проверяем наличие фактов
        visits_fact = await db.execute(
            select(Fact).where(
                Fact.study_id == document.study_id,
                Fact.fact_type == "soa",
                Fact.fact_key == "visits",
                Fact.created_from_doc_version_id == test_version_with_soa.id,
            )
        )
        visits_fact_obj = visits_fact.scalar_one_or_none()
        assert visits_fact_obj is not None
        assert visits_fact_obj.status in [FactStatus.EXTRACTED, FactStatus.NEEDS_REVIEW]
        assert "visits" in visits_fact_obj.value_json

        procedures_fact = await db.execute(
            select(Fact).where(
                Fact.study_id == document.study_id,
                Fact.fact_type == "soa",
                Fact.fact_key == "procedures",
                Fact.created_from_doc_version_id == test_version_with_soa.id,
            )
        )
        procedures_fact_obj = procedures_fact.scalar_one_or_none()
        assert procedures_fact_obj is not None
        assert "procedures" in procedures_fact_obj.value_json

        matrix_fact = await db.execute(
            select(Fact).where(
                Fact.study_id == document.study_id,
                Fact.fact_type == "soa",
                Fact.fact_key == "matrix",
                Fact.created_from_doc_version_id == test_version_with_soa.id,
            )
        )
        matrix_fact_obj = matrix_fact.scalar_one_or_none()
        assert matrix_fact_obj is not None
        assert "matrix" in matrix_fact_obj.value_json

    @pytest.mark.asyncio
    async def test_soa_extraction_creates_evidence(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест, что SoA extraction создаёт fact_evidence."""
        service = IngestionService(db)
        await service.ingest(test_version_with_soa.id)
        await db.commit()

        # Получаем document для study_id
        document = await db.get(Document, test_version_with_soa.document_id)
        assert document is not None

        # Получаем matrix fact
        matrix_fact = await db.execute(
            select(Fact).where(
                Fact.study_id == document.study_id,
                Fact.fact_type == "soa",
                Fact.fact_key == "matrix",
                Fact.created_from_doc_version_id == test_version_with_soa.id,
            )
        )
        matrix_fact_obj = matrix_fact.scalar_one_or_none()
        assert matrix_fact_obj is not None

        # Проверяем наличие evidence
        evidence_stmt = select(FactEvidence).where(FactEvidence.fact_id == matrix_fact_obj.id)
        evidence_result = await db.execute(evidence_stmt)
        evidence_list = evidence_result.scalars().all()

        assert len(evidence_list) > 0

        # Проверяем, что evidence ссылается на существующие anchors
        for evidence in evidence_list:
            assert evidence.evidence_role == EvidenceRole.PRIMARY
            anchor_stmt = select(Anchor).where(Anchor.anchor_id == evidence.anchor_id)
            anchor_result = await db.execute(anchor_stmt)
            anchor_obj = anchor_result.scalar_one_or_none()
            assert anchor_obj is not None
            assert anchor_obj.content_type == AnchorContentType.CELL

    @pytest.mark.asyncio
    async def test_soa_api_endpoint(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест API endpoint для получения SoA."""
        from app.api.v1.documents import get_soa

        # Сначала запускаем ingestion
        service = IngestionService(db)
        await service.ingest(test_version_with_soa.id)
        await db.commit()

        # Получаем SoA через API
        soa_result = await get_soa(test_version_with_soa.id, db=db)

        # Проверяем результат
        assert soa_result.table_index >= 0
        assert soa_result.section_path is not None
        assert len(soa_result.visits) > 0
        assert len(soa_result.procedures) > 0
        assert len(soa_result.matrix) > 0

        # Проверяем структуру visits
        for visit in soa_result.visits:
            assert visit.visit_id.startswith("V")
            assert visit.label is not None
            assert visit.anchor_id is not None

        # Проверяем структуру procedures
        for proc in soa_result.procedures:
            assert proc.proc_id.startswith("P")
            assert proc.label is not None
            assert proc.anchor_id is not None

        # Проверяем структуру matrix
        for entry in soa_result.matrix:
            assert entry.visit_id.startswith("V")
            assert entry.proc_id.startswith("P")
            assert entry.value is not None
            assert entry.anchor_id is not None

    @pytest.mark.asyncio
    async def test_soa_re_ingest_deletes_old_facts(
        self, db: AsyncSession, test_version_with_soa: DocumentVersion
    ):
        """Тест, что re-ingest удаляет старые facts и создаёт новые."""
        service = IngestionService(db)

        # Первый ingestion
        result1 = await service.ingest(test_version_with_soa.id)
        await db.commit()

        assert result1.soa_detected is True

        # Получаем document для study_id
        document = await db.get(Document, test_version_with_soa.document_id)
        assert document is not None

        # Подсчитываем факты после первого ingestion
        facts_stmt1 = select(Fact).where(
            Fact.study_id == document.study_id,
            Fact.fact_type == "soa",
            Fact.created_from_doc_version_id == test_version_with_soa.id,
        )
        facts_result1 = await db.execute(facts_stmt1)
        facts_count_1 = len(facts_result1.scalars().all())
        assert facts_count_1 > 0

        # Второй ingestion (re-ingest)
        result2 = await service.ingest(test_version_with_soa.id)
        await db.commit()

        assert result2.soa_detected is True

        # Подсчитываем факты после второго ingestion
        facts_stmt2 = select(Fact).where(
            Fact.study_id == document.study_id,
            Fact.fact_type == "soa",
            Fact.created_from_doc_version_id == test_version_with_soa.id,
        )
        facts_result2 = await db.execute(facts_stmt2)
        facts_count_2 = len(facts_result2.scalars().all())

        # Количество фактов должно быть таким же (пересозданы)
        assert facts_count_2 == facts_count_1

