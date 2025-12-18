"""
Тест: автодетект document_language на этапе upload, если язык не задан явно.
"""
from __future__ import annotations

import tempfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1.documents import upload_document_version
from app.db.enums import DocumentLifecycleStatus, DocumentLanguage, DocumentType, IngestionStatus, StudyStatus
from app.db.models.auth import Workspace
from app.db.models.studies import Document, DocumentVersion, Study


@pytest.fixture
async def test_workspace(db: AsyncSession) -> Workspace:
    ws = Workspace(name="Test Workspace")
    db.add(ws)
    await db.commit()
    await db.refresh(ws)
    return ws


@pytest.fixture
async def test_study(db: AsyncSession, test_workspace: Workspace) -> Study:
    s = Study(
        workspace_id=test_workspace.id,
        study_code="TEST-001",
        title="Test Study",
        status=StudyStatus.ACTIVE,
    )
    db.add(s)
    await db.commit()
    await db.refresh(s)
    return s


@pytest.fixture
async def test_document(db: AsyncSession, test_study: Study) -> Document:
    d = Document(
        workspace_id=test_study.workspace_id,
        study_id=test_study.id,
        doc_type=DocumentType.PROTOCOL,
        title="Test Protocol",
        lifecycle_status=DocumentLifecycleStatus.DRAFT,
    )
    db.add(d)
    await db.commit()
    await db.refresh(d)
    return d


def _make_ru_docx_bytes() -> bytes:
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = Path(tmp.name)
    tmp.close()
    doc = DocxDocument()
    doc.add_paragraph("ПРОТОКОЛ КЛИНИЧЕСКОГО ИССЛЕДОВАНИЯ")
    doc.add_paragraph("Цели исследования и дизайн.")
    doc.save(str(tmp_path))
    data = tmp_path.read_bytes()
    try:
        tmp_path.unlink(missing_ok=True)
    except Exception:
        pass
    return data


@pytest.mark.asyncio
async def test_upload_autodetects_ru_when_language_unknown(db: AsyncSession, test_document: Document):
    version = DocumentVersion(
        document_id=test_document.id,
        version_label="v1.0",
        source_file_uri=None,
        source_sha256=None,
        ingestion_status=IngestionStatus.UPLOADED,
        document_language=DocumentLanguage.UNKNOWN,
    )
    db.add(version)
    await db.commit()
    await db.refresh(version)

    content = _make_ru_docx_bytes()
    upload = UploadFile(filename="protocol_ru.docx", file=BytesIO(content))

    await upload_document_version(version.id, file=upload, db=db)

    await db.refresh(version)
    assert version.document_language == DocumentLanguage.RU
    assert version.ingestion_summary_json is not None
    assert version.ingestion_summary_json.get("document_language_source") == "auto_detect_upload"


