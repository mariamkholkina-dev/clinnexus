"""
Регресс-тест: после ingest должны создаваться chunks (Step 6) с non-null embedding.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.enums import DocumentLifecycleStatus, DocumentType, IngestionStatus, StudyStatus
from app.db.models.anchors import Chunk
from app.db.models.auth import Workspace
from app.db.models.studies import Document, DocumentVersion, Study
from app.services.ingestion import IngestionService


class TestChunkCreation:
    @pytest.fixture
    async def test_workspace(self, db: AsyncSession) -> Workspace:
        workspace = Workspace(name="Test Workspace")
        db.add(workspace)
        await db.commit()
        await db.refresh(workspace)
        return workspace

    @pytest.fixture
    async def test_study(self, db: AsyncSession, test_workspace: Workspace) -> Study:
        study = Study(
            workspace_id=test_workspace.id,
            study_code="TEST-001",
            title="Test Study",
            status=StudyStatus.ACTIVE,
        )
        db.add(study)
        await db.commit()
        await db.refresh(study)
        return study

    @pytest.fixture
    async def test_document(self, db: AsyncSession, test_study: Study) -> Document:
        document = Document(
            workspace_id=test_study.workspace_id,
            study_id=test_study.id,
            doc_type=DocumentType.PROTOCOL,
            title="Test Document",
            lifecycle_status=DocumentLifecycleStatus.DRAFT,
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        return document

    @pytest.fixture
    def golden_docx_file(self) -> Path:
        tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp_path = Path(tmp_file.name)
        tmp_file.close()

        doc = DocxDocument()
        doc.add_paragraph("Introduction", style="Heading 1")
        doc.add_paragraph("Protocol Version: 2.0")
        doc.add_paragraph("Amendment Date: 05 March 2021")
        doc.add_paragraph("Objectives", style="Heading 2")
        doc.add_paragraph("First objective", style="List Bullet")
        doc.add_paragraph("Second objective", style="List Bullet")
        doc.save(str(tmp_path))
        return tmp_path

    @pytest.fixture
    async def test_version_with_docx(
        self, db: AsyncSession, test_document: Document, golden_docx_file: Path
    ) -> DocumentVersion:
        abs_path = golden_docx_file.resolve()
        if abs_path.as_posix().startswith("/"):
            file_uri = f"file://{abs_path.as_posix()}"
        else:
            file_uri = f"file:///{abs_path.as_posix()}"

        version = DocumentVersion(
            document_id=test_document.id,
            version_label="v1.0",
            source_file_uri=file_uri,
            source_sha256="test_hash",
            ingestion_status=IngestionStatus.UPLOADED,
        )
        db.add(version)
        await db.commit()
        await db.refresh(version)
        return version

    @pytest.mark.asyncio
    async def test_ingest_creates_chunks_with_embedding(self, db: AsyncSession, test_version_with_docx: DocumentVersion):
        service = IngestionService(db)

        res1 = await service.ingest(test_version_with_docx.id)
        await db.commit()

        assert res1.chunks_created > 0

        chunks = (await db.execute(select(Chunk).where(Chunk.doc_version_id == test_version_with_docx.id))).scalars().all()
        assert len(chunks) > 0

        sample = chunks[0]
        assert sample.anchor_ids and len(sample.anchor_ids) > 0
        assert sample.embedding is not None

        # Re-ingest: rebuild идемпотентен, chunk_id стабилен
        chunk_ids_1 = {c.chunk_id for c in chunks}

        res2 = await service.ingest(test_version_with_docx.id)
        await db.commit()

        assert res2.chunks_created == res1.chunks_created
        chunks2 = (await db.execute(select(Chunk).where(Chunk.doc_version_id == test_version_with_docx.id))).scalars().all()
        assert len(chunks2) == len(chunks)
        assert {c.chunk_id for c in chunks2} == chunk_ids_1


