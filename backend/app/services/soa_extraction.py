"""Сервис для извлечения Schedule of Activities из DOCX документов."""
from __future__ import annotations

import re
import urllib.parse
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.logging import logger
from app.db.enums import AnchorContentType, DocumentLanguage, EvidenceRole, FactStatus
from app.db.models.anchors import Anchor
from app.db.models.facts import Fact, FactEvidence
from app.db.models.studies import Document as DocumentModel, DocumentVersion
from app.schemas.common import SoAResult, SoAVisit, SoAProcedure, SoAMatrixEntry, SoANote
from app.services.ingestion.docx_ingestor import normalize_text, get_text_hash, normalize_section_path, detect_text_language


@dataclass
class CellAnchorCreate:
    """Данные для создания cell anchor."""

    doc_version_id: UUID
    anchor_id: str
    section_path: str
    content_type: AnchorContentType
    ordinal: int
    text_raw: str
    text_norm: str
    text_hash: str
    location_json: dict[str, Any]
    language: DocumentLanguage = DocumentLanguage.UNKNOWN


@dataclass
class TableScore:
    """Оценка таблицы как потенциальной SoA."""

    table_index: int
    score: float
    section_path: str
    reason: str


# Версия SoAExtractionService (увеличивается при изменении логики извлечения SoA)
VERSION = "1.0.0"


class SoAExtractionService:
    """Сервис для извлечения Schedule of Activities из документа."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    def _iter_doc_body_blocks(self, doc: Document):
        """
        Итератор по блокам документа в исходном порядке: Paragraph / Table.
        Нужен, чтобы корректно брать ближайший контекст перед таблицей.
        """
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def _heading_level(self, para: Paragraph) -> int | None:
        if not para.style or not para.style.name:
            return None
        style = para.style.name
        if style.startswith("Heading") and style.replace("Heading", "").strip().isdigit():
            try:
                return int(style.replace("Heading", "").strip())
            except ValueError:
                return None
        return None

    def _get_context_for_table(
        self,
        doc: Document,
        table_index: int,
        *,
        max_prev_paras: int = 6,
        max_heading_stack: int = 4,
    ) -> tuple[str, list[str]]:
        """
        Возвращает (heading_context, prev_paragraphs) для таблицы с индексом table_index.
        heading_context: "H1 / H2 / ..."
        prev_paragraphs: тексты нескольких параграфов перед таблицей.
        """
        heading_stack: list[tuple[int, str]] = []
        prev_paras: list[str] = []
        cur_table_idx = -1

        for block in self._iter_doc_body_blocks(doc):
            if isinstance(block, Paragraph):
                lvl = self._heading_level(block)
                text_norm = normalize_text(block.text)
                if text_norm:
                    prev_paras.append(text_norm)
                    if len(prev_paras) > max_prev_paras:
                        prev_paras = prev_paras[-max_prev_paras:]
                if lvl and text_norm:
                    heading_stack = [h for h in heading_stack if h[0] < lvl]
                    heading_stack.append((lvl, text_norm))
                    if len(heading_stack) > max_heading_stack:
                        heading_stack = heading_stack[-max_heading_stack:]
            else:
                # Table
                cur_table_idx += 1
                if cur_table_idx == table_index:
                    break

        heading_context = " / ".join([h[1] for h in heading_stack]) if heading_stack else "ROOT"
        return heading_context, prev_paras

    def _soa_context_score(self, text: str) -> tuple[float, list[str]]:
        """
        Штрафы/бусты по ближайшему контексту (заголовки/параграфы).
        """
        t = (text or "").lower()
        reasons: list[str] = []
        score = 0.0

        # Позитивный контекст SoA
        pos = [
            "schedule of activities",
            "schedule of assessments",
            "schedule of procedures",
            "table of activities",
            "расписание процедур",
            "график процедур",
            "расписание мероприятий",
            "график мероприятий",
            "график проведения",  # Добавляем "график проведения процедур"
            "график проведения процедур",
            "график проведения исследований",
        ]
        # Негативный контекст (приложения/шкалы/анкеты)
        neg = [
            "appendix",
            "приложение",
            "scale",
            "шкала",
            "questionnaire",
            "опросник",
            "анкета",
            "form",
            "форма",
        ]

        for kw in neg:
            if kw in t:
                score -= 8.0
                reasons.append(f"-appendix_ctx:'{kw}'")
                # Если это appendix/scale контекст — не даём "позитивному" сигналу перебить штраф.
                return score, reasons

        for kw in pos:
            if kw in t:
                score += 6.0
                reasons.append(f"+soa_ctx:'{kw}'")
                break

        return score, reasons

    def _get_section_path_for_table(
        self,
        doc: Document,
        table_index: int,
        heading_stack: list[tuple[int, str]],
    ) -> str:
        """Определяет section_path для таблицы на основе ближайшего заголовка."""
        # Ищем заголовки перед таблицей
        # Для простоты используем текущий heading_stack из docx_ingestor логики
        # Но здесь нам нужно пройтись по параграфам до таблицы
        
        # Находим индекс таблицы в документе
        # В python-docx таблицы находятся в document.element.body
        # Проходимся по элементам до таблицы и собираем заголовки
        
        current_stack = heading_stack.copy()
        
        # Исторический API оставляем, но делаем корректный контекст по table_index.
        heading_context, _prev_paras = self._get_context_for_table(doc, table_index)
        if heading_context and heading_context != "ROOT":
            return normalize_section_path(heading_context.split(" / "))
        
        # Если стек пуст, используем ROOT
        if not current_stack:
            return "ROOT"
        
        path_parts = [h[1] for h in current_stack]
        return normalize_section_path(path_parts)

    def _score_table_for_soa(
        self,
        table: Table,
        table_index: int,
        doc: Document,
        heading_stack: list[tuple[int, str]],
    ) -> TableScore:
        """Оценивает таблицу на предмет того, является ли она SoA."""
        score = 0.0
        reasons: list[str] = []
        
        # Получаем ближайший контекст (заголовки/параграфы) для этой таблицы
        heading_context, prev_paras = self._get_context_for_table(doc, table_index)
        section_path = normalize_section_path(heading_context.split(" / ")) if heading_context else "ROOT"

        ctx_text = " ".join([heading_context] + prev_paras)
        ctx_delta, ctx_reasons = self._soa_context_score(ctx_text)
        score += ctx_delta
        reasons.extend(ctx_reasons)
        
        # Проверяем первую строку и первый столбец
        if len(table.rows) == 0:
            return TableScore(table_index, 0.0, section_path, "Пустая таблица")
        
        first_row_cells = [normalize_text(cell.text) for cell in table.rows[0].cells]
        first_col_cells = [normalize_text(table.rows[i].cells[0].text) if len(table.rows[i].cells) > 0 else "" 
                          for i in range(min(5, len(table.rows)))]
        
        # Логируем содержимое для отладки
        logger.debug(
            f"Таблица {table_index}: rows={len(table.rows)}, "
            f"first_row={first_row_cells[:5]}, first_col={first_col_cells[:5]}"
        )
        
        # Проверяем наличие "Visit", "Day", "Week", "Screening", "Baseline"
        # Расширяем поиск: проверяем первые 3 строки и первые 10 столбцов для лучшего распознавания
        visit_keywords = ["visit", "day", "week", "screening", "baseline", "визит", "день", "неделя", "цикл"]
        visit_keywords_found = []
        
        # Проверяем первую строку
        for cell_text in first_row_cells[:10]:
            cell_lower = cell_text.lower()
            for keyword in visit_keywords:
                if keyword in cell_lower:
                    if keyword not in visit_keywords_found:
                        score += 3.0
                        reasons.append(f"Найдено '{keyword}' в первой строке")
                        visit_keywords_found.append(keyword)
                    break
        
        # Проверяем первый столбец (первые 5 строк)
        for cell_text in first_col_cells[:5]:
            cell_lower = cell_text.lower()
            for keyword in visit_keywords:
                if keyword in cell_lower:
                    if keyword not in visit_keywords_found:
                        score += 3.0
                        reasons.append(f"Найдено '{keyword}' в первом столбце")
                        visit_keywords_found.append(keyword)
                    break
        
        # Дополнительно проверяем вторую и третью строки (для случаев, когда заголовок визитов не в первой строке)
        for row_idx in range(1, min(3, len(table.rows))):
            if len(table.rows[row_idx].cells) > 0:
                row_cells = [normalize_text(cell.text) for cell in table.rows[row_idx].cells[:10]]
                for cell_text in row_cells:
                    cell_lower = cell_text.lower()
                    for keyword in visit_keywords:
                        if keyword in cell_lower:
                            if keyword not in visit_keywords_found:
                                score += 3.0
                                reasons.append(f"Найдено '{keyword}' в строке {row_idx + 1}")
                                visit_keywords_found.append(keyword)
                            break
        
        # Проверяем наличие маркеров X, ✓, NA, —
        # Учитываем как латинскую "X", так и кириллическую "Х" (U+0425)
        marker_count = 0
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text).upper()
                # Нормализуем кириллическую "Х" к латинской "X" для распознавания
                cell_text_normalized = cell_text.replace("Х", "X")  # Кириллическая Х -> латинская X
                if cell_text_normalized in ["X", "✓", "NA", "—", "-", "YES", "NO"]:
                    marker_count += 1
        
        if marker_count > 5:
            score += 3.0
            reasons.append(f"Найдено {marker_count} маркеров (X, ✓, NA, etc.)")
        
        # Дополнительные критерии: структура таблицы
        # Если таблица имеет структуру матрицы (много строк и столбцов), это может быть SoA
        if len(table.rows) >= 3 and len(table.rows[0].cells) >= 3:
            score += 2.0
            reasons.append(f"Таблица имеет структуру матрицы ({len(table.rows)} строк, {len(table.rows[0].cells)} столбцов)")
        
        # Если таблица достаточно большая и содержит много пустых или маркированных ячеек
        total_cells = sum(len(row.cells) for row in table.rows)
        empty_or_marker_cells = 0
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text).upper()
                # Нормализуем кириллическую "Х" к латинской "X"
                cell_text_normalized = cell_text.replace("Х", "X")
                if not cell_text or cell_text_normalized in ["X", "✓", "NA", "—", "-", "YES", "NO", ""]:
                    empty_or_marker_cells += 1
        
        if total_cells > 0:
            marker_ratio = empty_or_marker_cells / total_cells
            if marker_ratio > 0.3:  # Более 30% ячеек пустые или маркированные
                score += 2.0
                reasons.append(f"Высокий процент маркированных/пустых ячеек ({marker_ratio:.1%})")
        
        # Штраф за таблицы с большим количеством чисел (lab ranges)
        number_count = 0
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text)
                # Проверяем, содержит ли ячейка много чисел и единиц измерения
                if re.search(r'\d+\.\d+|\d+\s*(mg|ml|g|kg|mmol|μmol)', cell_text, re.IGNORECASE):
                    number_count += 1
        
        if number_count > 10:
            score -= 5.0
            reasons.append(f"Таблица похожа на lab ranges ({number_count} числовых значений)")
        
        reason_str = "; ".join(reasons) if reasons else "Низкая оценка"
        
        # Детальное логирование для отладки
        logger.info(
            f"Оценка таблицы {table_index}: score={score:.1f}, "
            f"markers={marker_count}, numbers={number_count}, "
            f"heading_ctx={heading_context!r}, reasons={reason_str}"
        )
        
        return TableScore(table_index, score, section_path, reason_str)

    def _detect_soa_table(
        self,
        doc: Document,
        heading_stack: list[tuple[int, str]],
    ) -> TableScore | None:
        """Обнаруживает таблицу SoA в документе."""
        if len(doc.tables) == 0:
            return None
        
        scores: list[TableScore] = []
        
        for i, table in enumerate(doc.tables):
            score_result = self._score_table_for_soa(table, i, doc, heading_stack)
            scores.append(score_result)

        if not scores:
            return None

        scores_sorted = sorted(scores, key=lambda s: s.score, reverse=True)
        best_score = scores_sorted[0]

        # Логируем top-3 кандидата с контекстом
        top3 = scores_sorted[:3]
        for rank, s in enumerate(top3, start=1):
            heading_context, _ = self._get_context_for_table(doc, s.table_index)
            logger.info(
                "SoA кандидаты: "
                f"rank={rank}, table_index={s.table_index}, score={s.score:.1f}, "
                f"heading_ctx={heading_context!r}, section={s.section_path}, reason={s.reason}"
            )
        
        # Порог для принятия решения
        if best_score:
            logger.info(
                f"Лучшая таблица для SoA: индекс={best_score.table_index}, "
                f"score={best_score.score:.1f}, section={best_score.section_path}, "
                f"reason={best_score.reason}"
            )
            # Снижаем порог до 2.0 для более гибкого распознавания
            # Если таблица имеет подходящую структуру, принимаем её даже без явных ключевых слов
            if best_score.score >= 2.0:
                return best_score
            else:
                logger.info(f"Score {best_score.score:.1f} ниже порога 2.0, таблица не принята")
        else:
            logger.info("Не найдено ни одной таблицы в документе")
        
        return None

    def _determine_orientation(
        self,
        table: Table,
    ) -> tuple[bool, str]:
        """
        Определяет ориентацию таблицы.
        
        Returns:
            (is_rows_procedures, reason) - True если строки = процедуры, False если столбцы = процедуры
        """
        if len(table.rows) == 0:
            return (True, "Пустая таблица, используем ориентацию по умолчанию")
        
        first_row = [normalize_text(cell.text).lower() for cell in table.rows[0].cells]
        first_col = [normalize_text(table.rows[i].cells[0].text).lower() 
                    if len(table.rows[i].cells) > 0 else "" 
                    for i in range(min(5, len(table.rows)))]
        
        # Проверяем наличие visit keywords в первой строке
        visit_keywords = ["visit", "day", "week", "screening", "baseline", "визит"]
        first_row_visit_score = sum(1 for cell in first_row if any(kw in cell for kw in visit_keywords))
        first_col_visit_score = sum(1 for cell in first_col if any(kw in cell for kw in visit_keywords))
        
        # Проверяем плотность повторяющихся labels (процедуры обычно повторяются)
        # Процедуры обычно имеют более длинные и уникальные названия
        first_row_unique = len(set(first_row)) / max(len(first_row), 1)
        first_col_unique = len(set(first_col)) / max(len(first_col), 1)
        
        # Если в первой строке больше visit keywords, то строки = процедуры
        if first_row_visit_score > first_col_visit_score:
            return (True, f"Visit keywords найдены в первой строке ({first_row_visit_score})")
        elif first_col_visit_score > first_row_visit_score:
            return (False, f"Visit keywords найдены в первом столбце ({first_col_visit_score})")
        else:
            # Используем уникальность: процедуры обычно более уникальны
            if first_col_unique < first_row_unique:
                return (True, f"Первый столбец более уникален (процедуры)")
            else:
                return (False, f"Первая строка более уникальна (процедуры)")

    def _normalize_cell_value(self, value: str) -> str:
        """Нормализует значение ячейки."""
        value_norm = normalize_text(value).upper()
        
        if value_norm in ["X", "✓", "YES", "Y"]:
            return "X"
        elif value_norm in ["", "—", "-", "NA", "N/A", "NONE"]:
            return ""
        else:
            return normalize_text(value)

    def _extract_soa_from_table(
        self,
        table: Table,
        table_index: int,
        section_path: str,
        doc_version_id: UUID,
        ordinal_counters: dict[tuple[str, AnchorContentType], int],
    ) -> tuple[list[CellAnchorCreate], SoAResult]:
        """Извлекает SoA из таблицы и создаёт cell anchors."""
        cell_anchors: list[CellAnchorCreate] = []
        # Для устранения коллизий anchor_id в рамках одной версии
        anchor_id_counts: dict[str, int] = {}
        def _uniq_anchor_id(base: str) -> str:
            count = anchor_id_counts.get(base, 0) + 1
            anchor_id_counts[base] = count
            return base if count == 1 else f"{base}:v{count}"
        visits: list[SoAVisit] = []
        procedures: list[SoAProcedure] = []
        matrix: list[SoAMatrixEntry] = []
        notes: list[SoANote] = []
        
        if len(table.rows) == 0:
            return cell_anchors, SoAResult(
                table_index=table_index,
                section_path=section_path,
                visits=[],
                procedures=[],
                matrix=[],
                notes=[],
                confidence=0.0,
                warnings=["Таблица пуста"],
            )
        
        # Определяем ориентацию
        is_rows_procedures, orientation_reason = self._determine_orientation(table)
        
        # Извлекаем visits и procedures в зависимости от ориентации
        if is_rows_procedures:
            # Первая строка = visits, первый столбец = procedures
            header_row = table.rows[0]
            visit_labels = [normalize_text(cell.text) for cell in header_row.cells[1:]]  # Пропускаем первую ячейку
            
            procedure_labels: list[str] = []
            for i in range(1, len(table.rows)):
                if len(table.rows[i].cells) > 0:
                    proc_label = normalize_text(table.rows[i].cells[0].text)
                    if proc_label:
                        procedure_labels.append(proc_label)
        else:
            # Первый столбец = visits, первая строка = procedures
            visit_labels = []
            for i in range(1, len(table.rows)):
                if len(table.rows[i].cells) > 0:
                    visit_label = normalize_text(table.rows[i].cells[0].text)
                    if visit_label:
                        visit_labels.append(visit_label)
            
            header_row = table.rows[0]
            procedure_labels = [normalize_text(cell.text) for cell in header_row.cells[1:]]
        
        logger.info(
            f"Извлечено labels: visits={len(visit_labels)}, procedures={len(procedure_labels)}, "
            f"orientation={orientation_reason}"
        )
        
        # Создаём visits с anchors
        visit_anchor_ids: dict[int, str] = {}
        for idx, label in enumerate(visit_labels):
            if not label:
                continue
            
            visit_id = f"V{idx + 1}"
            
            # Создаём anchor для header cell
            if is_rows_procedures:
                # Header в первой строке, столбец idx+1
                if len(table.rows[0].cells) > idx + 1:
                    header_cell = table.rows[0].cells[idx + 1]
                else:
                    continue
            else:
                # Header в первом столбце, строка idx+1
                if len(table.rows) > idx + 1 and len(table.rows[idx + 1].cells) > 0:
                    header_cell = table.rows[idx + 1].cells[0]
                else:
                    continue
            
            text_raw = header_cell.text
            text_norm = normalize_text(text_raw)
            text_hash = get_text_hash(text_norm)
            
            # Определяем язык для cell anchor (RU → ru, EN → en, иначе unknown)
            cell_language = detect_text_language(text_norm)
            
            key = (section_path, AnchorContentType.CELL)
            ordinal = ordinal_counters.get(key, 0) + 1
            ordinal_counters[key] = ordinal
            
            # Определяем row_idx и col_idx (нужно для anchor_id)
            if is_rows_procedures:
                row_idx = 0
                col_idx = idx + 1
            else:
                row_idx = idx + 1
                col_idx = 0
            
            # Формируем anchor_id для header cell:
            # {doc_version_id}:cell:{get_text_hash(text_norm + row_idx + col_idx)[:16]}
            cell_hash_input = f"{text_norm}:{row_idx}:{col_idx}"
            cell_hash = get_text_hash(cell_hash_input)[:16]
            base_anchor_id = f"{doc_version_id}:cell:{cell_hash}"
            anchor_id = _uniq_anchor_id(base_anchor_id)
            visit_anchor_ids[idx] = anchor_id
            
            location_json = {
                "table_id": table_index,
                "table_index": table_index,
                "row_idx": row_idx,
                "col_idx": col_idx,
                "is_header": True,
                "header_path": {"row_headers": [], "col_headers": []},
                "para_index_hint": None,
            }
            
            cell_anchor = CellAnchorCreate(
                doc_version_id=doc_version_id,
                anchor_id=anchor_id,
                section_path=section_path,
                content_type=AnchorContentType.CELL,
                ordinal=ordinal,
                text_raw=text_raw,
                text_norm=text_norm,
                text_hash=text_hash,
                location_json=location_json,
                language=cell_language,
            )
            cell_anchors.append(cell_anchor)
            
            visits.append(SoAVisit(
                visit_id=visit_id,
                label=label,
                day=None,  # TODO: извлечь из label если возможно
                anchor_id=anchor_id,
            ))
        
        # Создаём procedures с anchors
        proc_anchor_ids: dict[int, str] = {}
        for idx, label in enumerate(procedure_labels):
            if not label:
                continue
            
            proc_id = f"P{idx + 1:03d}"
            
            # Создаём anchor для header cell
            if is_rows_procedures:
                # Header в первом столбце, строка idx+1
                if len(table.rows) > idx + 1 and len(table.rows[idx + 1].cells) > 0:
                    header_cell = table.rows[idx + 1].cells[0]
                else:
                    continue
            else:
                # Header в первой строке, столбец idx+1
                if len(table.rows[0].cells) > idx + 1:
                    header_cell = table.rows[0].cells[idx + 1]
                else:
                    continue
            
            text_raw = header_cell.text
            text_norm = normalize_text(text_raw)
            text_hash = get_text_hash(text_norm)
            
            # Определяем язык для cell anchor (RU → ru, EN → en, иначе unknown)
            cell_language = detect_text_language(text_norm)
            
            key = (section_path, AnchorContentType.CELL)
            ordinal = ordinal_counters.get(key, 0) + 1
            ordinal_counters[key] = ordinal
            
            # Определяем row_idx и col_idx (нужно для anchor_id)
            if is_rows_procedures:
                row_idx = idx + 1
                col_idx = 0
            else:
                row_idx = 0
                col_idx = idx + 1
            
            # Формируем anchor_id для procedure header cell:
            # {doc_version_id}:cell:{get_text_hash(text_norm + row_idx + col_idx)[:16]}
            cell_hash_input = f"{text_norm}:{row_idx}:{col_idx}"
            cell_hash = get_text_hash(cell_hash_input)[:16]
            base_anchor_id = f"{doc_version_id}:cell:{cell_hash}"
            anchor_id = _uniq_anchor_id(base_anchor_id)
            proc_anchor_ids[idx] = anchor_id
            
            location_json = {
                "table_id": table_index,
                "table_index": table_index,
                "row_idx": row_idx,
                "col_idx": col_idx,
                "is_header": True,
                "header_path": {"row_headers": [], "col_headers": []},
                "para_index_hint": None,
            }
            
            cell_anchor = CellAnchorCreate(
                doc_version_id=doc_version_id,
                anchor_id=anchor_id,
                section_path=section_path,
                content_type=AnchorContentType.CELL,
                ordinal=ordinal,
                text_raw=text_raw,
                text_norm=text_norm,
                text_hash=text_hash,
                location_json=location_json,
                language=cell_language,
            )
            cell_anchors.append(cell_anchor)
            
            procedures.append(SoAProcedure(
                proc_id=proc_id,
                label=label,
                category=None,  # TODO: извлечь категорию если возможно
                anchor_id=anchor_id,
            ))
        
        # Создаём matrix entries с anchors
        for proc_idx, proc_id in enumerate([p.proc_id for p in procedures]):
            for visit_idx, visit_id in enumerate([v.visit_id for v in visits]):
                # Определяем координаты ячейки
                if is_rows_procedures:
                    row_idx = proc_idx + 1
                    col_idx = visit_idx + 1
                else:
                    row_idx = visit_idx + 1
                    col_idx = proc_idx + 1
                
                # Получаем значение ячейки
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    value_raw = cell.text
                    value_norm = self._normalize_cell_value(value_raw)
                else:
                    value_norm = ""
                
                # Создаём anchor для ячейки (даже если пустая, для traceability)
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    text_raw = cell.text
                    text_norm = normalize_text(text_raw)
                    text_hash = get_text_hash(text_norm)
                    
                    # Определяем язык для cell anchor (RU → ru, EN → en, иначе unknown)
                    cell_language = detect_text_language(text_norm)
                    
                    key = (section_path, AnchorContentType.CELL)
                    ordinal = ordinal_counters.get(key, 0) + 1
                    ordinal_counters[key] = ordinal
                    
                    # Определяем header_path
                    row_headers = [procedures[proc_idx].label] if proc_idx < len(procedures) else []
                    col_headers = [visits[visit_idx].label] if visit_idx < len(visits) else []
                    
                    # Формируем anchor_id для body cell:
                    # {doc_version_id}:cell:{get_text_hash(text_norm + row_idx + col_idx)[:16]}
                    cell_hash_input = f"{value_norm}:{row_idx}:{col_idx}"
                    cell_hash = get_text_hash(cell_hash_input)[:16]
                    base_anchor_id = f"{doc_version_id}:cell:{cell_hash}"
                    anchor_id = _uniq_anchor_id(base_anchor_id)
                    
                    location_json = {
                        "table_id": table_index,
                        "table_index": table_index,
                        "row_idx": row_idx,
                        "col_idx": col_idx,
                        "is_header": False,
                        "header_path": {
                            "row_headers": row_headers,
                            "col_headers": col_headers,
                        },
                        "para_index_hint": None,
                    }
                    
                    cell_anchor = CellAnchorCreate(
                        doc_version_id=doc_version_id,
                        anchor_id=anchor_id,
                        section_path=section_path,
                        content_type=AnchorContentType.CELL,
                        ordinal=ordinal,
                        text_raw=text_raw,
                        text_norm=text_norm,
                        text_hash=text_hash,
                        location_json=location_json,
                        language=cell_language,
                    )
                    cell_anchors.append(cell_anchor)
                    
                    # Добавляем в matrix только non-empty значения
                    if value_norm:
                        matrix.append(SoAMatrixEntry(
                            visit_id=visit_id,
                            proc_id=proc_id,
                            value=value_norm,
                            anchor_id=anchor_id,
                        ))
        
        confidence = 0.7 if len(visits) > 0 and len(procedures) > 0 else 0.3
        warnings: list[str] = []
        if not visits:
            warnings.append("Не найдены визиты")
        if not procedures:
            warnings.append("Не найдены процедуры")
        
        logger.info(
            f"_extract_soa_from_table завершён: cell_anchors={len(cell_anchors)}, "
            f"visits={len(visits)}, procedures={len(procedures)}, matrix={len(matrix)}"
        )
        
        return cell_anchors, SoAResult(
            table_index=table_index,
            section_path=section_path,
            visits=visits,
            procedures=procedures,
            matrix=matrix,
            notes=notes,
            confidence=confidence,
            warnings=warnings,
        )

    async def extract_soa(
        self,
        doc_version_id: UUID,
    ) -> tuple[list[CellAnchorCreate], SoAResult | None]:
        """
        Извлекает Schedule of Activities из версии документа.

        Returns:
            (cell_anchors, soa_result) - список cell anchors и результат извлечения SoA
        """
        logger.info(f"Извлечение SoA из документа {doc_version_id}")

        # Получаем версию документа
        doc_version = await self.db.get(DocumentVersion, doc_version_id)
        if not doc_version:
            raise ValueError(f"DocumentVersion {doc_version_id} не найден")
        
        # Получаем документ для получения study_id
        document = await self.db.get(DocumentModel, doc_version.document_id)
        if not document:
            raise ValueError(f"Document {doc_version.document_id} не найден")
        
        # Проверяем наличие файла
        if not doc_version.source_file_uri:
            logger.warning(f"DocumentVersion {doc_version_id} не имеет source_file_uri")
            return [], None
        
        # Загружаем DOCX документ
        # Преобразуем URI в локальный путь
        uri = doc_version.source_file_uri
        if uri.startswith("file://"):
            parsed = urllib.parse.urlparse(uri)
            path = urllib.parse.unquote(parsed.path)
            # На Windows file:///C:/path становится /C:/path, убираем ведущий /
            if path.startswith("/") and len(path) > 3 and path[2] == ":":
                path = path[1:]
            file_path = Path(path)
        else:
            file_path = Path(uri)
        
        if not file_path.exists():
            logger.warning(f"Файл не найден: {file_path}")
            return [], None
        
        if file_path.suffix.lower() != ".docx":
            logger.warning(f"Файл не является DOCX: {file_path}")
            return [], None
        
        doc = Document(str(file_path))
        
        # Собираем heading_stack из параграфов (для определения section_path)
        heading_stack: list[tuple[int, str]] = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            if style_name.startswith('Heading') and style_name.replace('Heading', '').strip().isdigit():
                level_str = style_name.replace('Heading', '').strip()
                try:
                    level = int(level_str)
                    text_norm = normalize_text(para.text)
                    if text_norm:
                        heading_stack = [h for h in heading_stack if h[0] < level]
                        heading_stack.append((level, text_norm))
                except ValueError:
                    pass
        
        # Обнаруживаем SoA таблицу
        soa_table_score = self._detect_soa_table(doc, heading_stack)
        
        if not soa_table_score:
            logger.info(f"SoA таблица не найдена в документе {doc_version_id}")
            return [], None
        
        logger.info(
            f"Найдена SoA таблица (индекс {soa_table_score.table_index}, "
            f"score={soa_table_score.score:.1f}, section={soa_table_score.section_path})"
        )
        
        # Извлекаем SoA из таблицы
        table = doc.tables[soa_table_score.table_index]
        ordinal_counters: dict[tuple[str, AnchorContentType], int] = {}
        
        cell_anchors, soa_result = self._extract_soa_from_table(
            table,
            soa_table_score.table_index,
            soa_table_score.section_path,
            doc_version_id,
            ordinal_counters,
        )
        
        # Обновляем confidence на основе score
        soa_result.confidence = min(0.9, 0.5 + (soa_table_score.score / 20.0))
        
        logger.info(
            f"SoA извлечён: {len(soa_result.visits)} visits, {len(soa_result.procedures)} procedures, "
            f"{len(soa_result.matrix)} matrix entries, {len(cell_anchors)} cell anchors"
        )
        
        return cell_anchors, soa_result
