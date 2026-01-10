"""Сервис сборки финального документа из сгенерированных секций."""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.core.logging import logger

from uuid import UUID

if TYPE_CHECKING:
    from app.db.models.generation import GeneratedTargetSection, GenerationRun
else:
    from app.db.models.generation import GenerationRun


# Регулярное выражение для поиска маркеров якорей [anchor:xyz]
ANCHOR_MARKER_PATTERN = re.compile(r'\[anchor:[^\]]+\]', re.IGNORECASE)


def _remove_anchor_markers(text: str) -> str:
    """Удаляет маркеры якорей из текста (для MVP просто удаляем)."""
    return ANCHOR_MARKER_PATTERN.sub('', text).strip()


def _get_section_display_name(section_key: str) -> str:
    """Получает отображаемое имя секции по её ключу."""
    # Простая маппинг для MVP (можно улучшить позже)
    display_names: dict[str, str] = {
        "overview": "Обзор",
        "design": "Дизайн исследования",
        "ip": "Исследуемый препарат",
        "statistics": "Статистика",
        "safety": "Безопасность",
        "endpoints": "Конечные точки",
        "population": "Популяция",
        "procedures": "Процедуры",
        "data_management": "Управление данными",
        "ethics": "Этика",
        "admin": "Административные вопросы",
        "appendix": "Приложение",
    }
    return display_names.get(section_key, section_key.replace("_", " ").title())


def assemble_document(
    sections: list[GeneratedTargetSection],
    output_path: Path,
    *,
    section_ordering: list[str] | None = None,
    generation_runs: dict[UUID, GenerationRun] | None = None,
) -> None:
    """
    Собирает финальный документ из списка сгенерированных секций.
    
    Args:
        sections: Список сгенерированных секций (GeneratedTargetSection)
        output_path: Путь для сохранения собранного документа
        section_ordering: Опциональный порядок секций (по ключам).
                         Если None, используется порядок из CANONICAL_SECTION_KEYS.
        generation_runs: Опциональный словарь generation_run_id -> GenerationRun.
                        Если None, попытается использовать section.generation_run (если relationship настроен).
    
    Raises:
        ValueError: Если список секций пуст
        IOError: Если не удалось сохранить файл
    """
    if not sections:
        raise ValueError("Список секций не может быть пустым")
    
    # Импортируем здесь, чтобы избежать циклических зависимостей
    from app.core.section_standardization import CANONICAL_SECTION_KEYS
    
    # Создаём пустой документ
    doc = Document()
    
    # Устанавливаем базовые стили (опционально)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Получаем ключи секций из generation_runs
    # Для этого нужно получить target_section из связанного GenerationRun
    section_data: list[tuple[str, str]] = []  # (target_section, content_text)
    
    for section in sections:
        # Получаем target_section из связанного GenerationRun
        generation_run: GenerationRun | None = None
        
        if generation_runs:
            # Используем переданный словарь
            generation_run = generation_runs.get(section.generation_run_id)
        else:
            # Пытаемся использовать relationship (если настроен)
            generation_run = getattr(section, "generation_run", None)
        
        if not generation_run:
            logger.warning(
                f"GenerationRun не найден для секции {section.id} "
                f"(generation_run_id={section.generation_run_id}). "
                "Пропускаем секцию."
            )
            continue
        
        target_section = generation_run.target_section
        content_text = _remove_anchor_markers(section.content_text)
        
        if not content_text.strip():
            logger.warning(
                f"Секция {target_section} (ID: {section.id}) пуста. Пропускаем."
            )
            continue
        
        section_data.append((target_section, content_text))
    
    if not section_data:
        raise ValueError("Нет секций с содержимым для сборки документа")
    
    # Определяем порядок сортировки
    if section_ordering is None:
        section_ordering = CANONICAL_SECTION_KEYS.copy()
    
    # Функция для получения индекса секции в порядке сортировки
    def get_section_order(section_key: str) -> int:
        """Возвращает индекс секции в порядке сортировки."""
        try:
            return section_ordering.index(section_key)
        except ValueError:
            # Если секция не в списке, добавляем в конец
            return len(section_ordering) + 1000
    
    # Сортируем секции по порядку
    section_data.sort(key=lambda x: get_section_order(x[0]))
    
    # Добавляем секции в документ
    for target_section, content_text in section_data:
        # Добавляем заголовок секции (Heading 1)
        heading = doc.add_heading(_get_section_display_name(target_section), level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Добавляем содержимое секции
        # Разбиваем текст на параграфы (по двойным переводам строк)
        paragraphs = content_text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Если параграф содержит переводы строк внутри, разбиваем на абзацы
            if '\n' in para_text:
                lines = para_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
            else:
                doc.add_paragraph(para_text)
        
        # Добавляем пустую строку между секциями
        doc.add_paragraph("")
    
    # Сохраняем документ
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(
            f"Документ успешно собран: {output_path} "
            f"(секций: {len(section_data)})"
        )
    except Exception as e:
        logger.error(f"Ошибка при сохранении документа {output_path}: {e}")
        raise IOError(f"Не удалось сохранить документ: {e}") from e

