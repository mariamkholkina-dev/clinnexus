"""
Регрессионные тесты для завершения MVP шагов 1–6 (без chunks, которые уже реализованы).
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1.documents import start_ingestion
from app.db.enums import (
    AnchorContentType,
    DocumentLifecycleStatus,
    DocumentType,
    IngestionStatus,
    StudyStatus,
)
from app.db.models.anchors import Anchor
from app.db.models.auth import Workspace
from app.db.models.facts import FactEvidence, Fact
from app.db.models.studies import Document, DocumentVersion, Study
from app.services.ingestion import IngestionService


@pytest.fixture
async def test_workspace(db: AsyncSession) -> Workspace:
    ws = Workspace(name="Test Workspace")
    db.add(ws)
    await db.commit()
    await db.refresh(ws)
    return ws


@pytest.fixture
async def test_study(db: AsyncSession, test_workspace: Workspace) -> Study:
    study = Study(
        workspace_id=test_workspace.id,
        study_code="TEST-001",
        title="Test Study",
        status=StudyStatus.ACTIVE,
    )
    db.add(study)
    await db.commit()
    await db.refresh(study)
    return study


@pytest.fixture
async def test_document(db: AsyncSession, test_study: Study) -> Document:
    doc = Document(
        workspace_id=test_study.workspace_id,
        study_id=test_study.id,
        doc_type=DocumentType.PROTOCOL,
        title="Test Protocol",
        lifecycle_status=DocumentLifecycleStatus.DRAFT,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


def _to_file_uri(path: Path) -> str:
    abs_path = path.resolve()
    if abs_path.as_posix().startswith("/"):
        return f"file://{abs_path.as_posix()}"
    return f"file:///{abs_path.as_posix()}"


@pytest.mark.asyncio
async def test_ingestion_summary_schema_present_on_success_like_flow(
    db: AsyncSession, test_document: Document, tmp_path: Path
):
    # PDF-заглушка: ingestion отработает без исключения, но поставит needs_review.
    pdf_path = tmp_path / "test.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%test\n")

    version = DocumentVersion(
        document_id=test_document.id,
        version_label="v1.0",
        source_file_uri=_to_file_uri(pdf_path),
        source_sha256="test_sha256",
        ingestion_status=IngestionStatus.UPLOADED,
        ingestion_summary_json={"filename": "test.pdf", "size_bytes": len(pdf_path.read_bytes())},
    )
    db.add(version)
    await db.commit()
    await db.refresh(version)

    res = await start_ingestion(version.id, force=False, db=db)
    assert res["status"] in (IngestionStatus.READY.value, IngestionStatus.NEEDS_REVIEW.value)

    await db.refresh(version)
    assert version.ingestion_summary_json is not None
    summary = version.ingestion_summary_json

    # Стабильная схема (всегда присутствует)
    for key in (
        "anchors_created",
        "soa_found",
        "soa_facts_written",
        "chunks_created",
        "mapping_status",
        "warnings",
        "errors",
    ):
        assert key in summary

    assert isinstance(summary["warnings"], list)
    assert isinstance(summary["errors"], list)
    assert summary.get("source_sha256") == "test_sha256"


@pytest.mark.asyncio
async def test_ingestion_summary_schema_present_on_failure(
    db: AsyncSession, test_document: Document
):
    # Несуществующий файл → FileNotFoundError внутри ingestion → статус FAILED + summary с errors
    fake_path = Path("C:/this/does/not/exist/protocol_v1.docx")
    version = DocumentVersion(
        document_id=test_document.id,
        version_label="v1.0",
        source_file_uri=_to_file_uri(fake_path),
        source_sha256="test_sha256",
        ingestion_status=IngestionStatus.UPLOADED,
    )
    db.add(version)
    await db.commit()
    await db.refresh(version)

    with pytest.raises(Exception):
        await start_ingestion(version.id, force=False, db=db)

    await db.refresh(version)
    assert version.ingestion_status == IngestionStatus.FAILED
    summary = version.ingestion_summary_json
    assert summary is not None
    for key in (
        "anchors_created",
        "soa_found",
        "soa_facts_written",
        "chunks_created",
        "mapping_status",
        "warnings",
        "errors",
    ):
        assert key in summary
    assert isinstance(summary["errors"], list)
    assert len(summary["errors"]) >= 1


@pytest.mark.asyncio
async def test_docx_anchors_include_hdr_p_li_and_fn_attempt(
    db: AsyncSession, test_document: Document
):
    # Создаём DOCX с заголовком, параграфом и list item.
    tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = Path(tmp_file.name)
    tmp_file.close()

    doc = DocxDocument()
    doc.add_paragraph("Schedule of Activities", style="Heading 1")
    doc.add_paragraph("Обычный параграф с текстом.")
    doc.add_paragraph("Пункт списка", style="List Bullet")
    doc.save(str(tmp_path))

    version = DocumentVersion(
        document_id=test_document.id,
        version_label="v1.0",
        source_file_uri=_to_file_uri(tmp_path),
        source_sha256="test_hash",
        ingestion_status=IngestionStatus.UPLOADED,
    )
    db.add(version)
    await db.commit()
    await db.refresh(version)

    # Запускаем через API-оркестратор, чтобы гарантированно заполнить ingestion_summary_json (включая warnings).
    await start_ingestion(version.id, force=False, db=db)

    anchors_res = await db.execute(select(Anchor).where(Anchor.doc_version_id == version.id))
    anchors = anchors_res.scalars().all()
    assert len(anchors) > 0

    types = {a.content_type for a in anchors}
    assert AnchorContentType.HDR in types
    assert AnchorContentType.P in types
    assert AnchorContentType.LI in types

    # Footnotes: либо реально созданы FN якоря, либо ingestion предупредил, что недоступны.
    has_fn = AnchorContentType.FN in types
    if not has_fn:
        # fallback: проверяем warning из DocxIngestor через summary в ingestion (он кладётся в docx_summary)
        await db.refresh(version)
        summary = version.ingestion_summary_json or {}
        warnings = summary.get("warnings") or []
        assert any("footnote" in str(w).lower() for w in warnings) or any("сноск" in str(w).lower() for w in warnings)


@pytest.mark.asyncio
async def test_soa_cell_anchors_and_evidence_are_renderable(
    db: AsyncSession, test_document: Document
):
    # Минимальный golden-подобный SoA DOCX: таблица с visits/procedures/matrix.
    tmp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = Path(tmp_file.name)
    tmp_file.close()

    doc = DocxDocument()
    doc.add_paragraph("Schedule of Activities", style="Heading 1")
    table = doc.add_table(rows=4, cols=4)
    table.rows[0].cells[0].text = "Procedure"
    table.rows[0].cells[1].text = "Screening"
    table.rows[0].cells[2].text = "Baseline"
    table.rows[0].cells[3].text = "Week 4"
    table.rows[1].cells[0].text = "Informed consent"
    table.rows[1].cells[1].text = "X"
    table.rows[1].cells[2].text = "X"
    table.rows[1].cells[3].text = ""
    table.rows[2].cells[0].text = "Vitals"
    table.rows[2].cells[1].text = "X"
    table.rows[2].cells[2].text = "X"
    table.rows[2].cells[3].text = "X"
    table.rows[3].cells[0].text = "ECG"
    table.rows[3].cells[1].text = ""
    table.rows[3].cells[2].text = "X"
    table.rows[3].cells[3].text = ""
    doc.save(str(tmp_path))

    version = DocumentVersion(
        document_id=test_document.id,
        version_label="v1.0",
        source_file_uri=_to_file_uri(tmp_path),
        source_sha256="test_hash",
        ingestion_status=IngestionStatus.UPLOADED,
    )
    db.add(version)
    await db.commit()
    await db.refresh(version)

    service = IngestionService(db)
    result = await service.ingest(version.id)
    await db.commit()

    assert result.soa_detected is True

    # Проверяем, что evidence ссылается на существующие cell anchors (рендеримо).
    facts_res = await db.execute(
        select(Fact).where(
            Fact.created_from_doc_version_id == version.id,
            Fact.fact_type == "soa",
            Fact.fact_key.in_(("visits", "procedures", "matrix")),
        )
    )
    facts = facts_res.scalars().all()
    assert len(facts) >= 1

    for f in facts:
        ev_res = await db.execute(select(FactEvidence).where(FactEvidence.fact_id == f.id))
        ev_list = ev_res.scalars().all()
        assert len(ev_list) > 0
        for ev in ev_list:
            a_res = await db.execute(select(Anchor).where(Anchor.anchor_id == ev.anchor_id))
            a = a_res.scalar_one_or_none()
            assert a is not None
            assert a.content_type == AnchorContentType.CELL
            assert isinstance(a.location_json, dict)
            # Минимальные поля для таблиц
            assert "table_id" in a.location_json
            assert "row_idx" in a.location_json
            assert "col_idx" in a.location_json


