"""Парсер DOCX документов для создания anchors."""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

from app.db.enums import AnchorContentType, DocumentLanguage
from app.services.ingestion.heading_detector import HeadingDetector, HeadingHit, DocStats
from app.services.source_zone_classifier import get_classifier

# Регулярные выражения для детекции языка
_CYR_RE = re.compile(r"[А-Яа-яЁё]")
_LAT_RE = re.compile(r"[A-Za-z]")


def normalize_text(text: str) -> str:
    """
    Нормализует текст для создания стабильного хеша.
    
    Правила нормализации:
    - trim (удаление пробелов по краям)
    - последовательности whitespace → один пробел
    - сохранение цифр/пунктуации
    - пустые строки возвращают пустую строку
    
    Args:
        text: Исходный текст
        
    Returns:
        Нормализованный текст
    """
    if not text:
        return ""
    
    # Заменяем последовательности whitespace на один пробел
    normalized = re.sub(r'\s+', ' ', text)
    # Удаляем пробелы по краям
    normalized = normalized.strip()
    
    return normalized


def get_text_hash(text_norm: str) -> str:
    """
    Вычисляет SHA256 хеш нормализованного текста.
    
    Args:
        text_norm: Нормализованный текст
        
    Returns:
        Hex-строка хеша (64 символа)
    """
    return hashlib.sha256(text_norm.encode('utf-8')).hexdigest()


def detect_text_language(text: str) -> DocumentLanguage:
    """
    Быстрая локальная детекция языка из текста (regex на кириллицу/латиницу).
    Используется для определения языка конкретного anchor/chunk.
    """
    if not text:
        return DocumentLanguage.UNKNOWN
    
    cyr_count = len(_CYR_RE.findall(text))
    lat_count = len(_LAT_RE.findall(text))
    total_letters = cyr_count + lat_count
    
    if total_letters == 0:
        return DocumentLanguage.UNKNOWN
    
    cyr_ratio = cyr_count / total_letters
    
    if cyr_ratio >= 0.7:
        return DocumentLanguage.RU
    elif cyr_ratio <= 0.3:
        return DocumentLanguage.EN
    else:
        # Смешанный, если достаточно обеих букв
        return DocumentLanguage.MIXED if cyr_count >= 10 and lat_count >= 10 else DocumentLanguage.UNKNOWN


# Константа для титульной страницы (frontmatter)
FRONTMATTER_SECTION = "__FRONTMATTER__"


def normalize_section_path(path_parts: list[str]) -> str:
    """
    Нормализует путь секции: trim + collapse spaces.
    
    Args:
        path_parts: Список частей пути (заголовков)
        
    Returns:
        Нормализованный путь (например "H1/H2/H3" или "ROOT" или "__FRONTMATTER__")
    """
    if not path_parts:
        return "ROOT"
    
    # Нормализуем каждую часть (trim + collapse spaces)
    normalized_parts = []
    for part in path_parts:
        normalized = re.sub(r'\s+', ' ', part.strip())
        if normalized:
            normalized_parts.append(normalized)
    
    if not normalized_parts:
        return "ROOT"
    
    return "/".join(normalized_parts)


def is_list_item(paragraph: Paragraph) -> bool:
    """
    Проверяет, является ли параграф элементом списка.
    
    Проверяет:
    1. Наличие numPr (numbering properties) в XML представлении параграфа
    2. Стиль параграфа (если начинается с "List")
    
    Args:
        paragraph: Параграф из python-docx
        
    Returns:
        True, если параграф является элементом списка
    """
    # Проверяем стиль параграфа
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("List"):
        return True
    
    # Проверяем наличие numbering properties
    try:
        p_element: CT_P = paragraph._p
        if p_element.pPr is not None and p_element.pPr.numPr is not None:
            return True
    except (AttributeError, TypeError):
        pass
    
    return False


@dataclass
class AnchorCreate:
    """Данные для создания anchor в БД."""
    
    doc_version_id: UUID
    anchor_id: str
    section_path: str
    content_type: AnchorContentType
    ordinal: int
    text_raw: str
    text_norm: str
    text_hash: str
    location_json: dict[str, Any]
    source_zone: str = "unknown"
    language: DocumentLanguage = DocumentLanguage.UNKNOWN


# Версия DocxIngestor (увеличивается при изменении логики парсинга)
VERSION = "1.0.0"


@dataclass
class DocxIngestResult:
    """Результат парсинга DOCX документа."""
    
    anchors: list[AnchorCreate]
    summary: dict[str, Any]
    warnings: list[str]


class DocxIngestor:
    """Парсер DOCX документов для создания anchors."""
    
    def __init__(self) -> None:
        """Инициализация DocxIngestor с классификатором source_zone."""
        self.zone_classifier = get_classifier()
    
    def ingest(self, file_path: str | Path, doc_version_id: UUID, document_language: DocumentLanguage) -> DocxIngestResult:
        """
        Парсит DOCX документ и создаёт anchors.
        
        Args:
            file_path: Путь к DOCX файлу
            doc_version_id: UUID версии документа
            document_language: Язык документа
            
        Returns:
            DocxIngestResult с anchors, summary и warnings
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        # Загружаем документ
        doc = Document(str(file_path))
        
        # Вычисляем статистику документа для visual fallback
        doc_stats = HeadingDetector.compute_doc_stats(list(doc.paragraphs))
        
        # Первый проход: детект только style/outline/numbering
        detector = HeadingDetector(enable_visual_fallback=False)
        detector.set_doc_stats(doc_stats)
        
        # Счётчик отклонений для numbering детекции
        rejection_counter: dict[str, int] = {}
        detector._rejection_counter = rejection_counter
        
        # Собираем hits для всех параграфов
        paragraph_hits: list[tuple[Paragraph, HeadingHit]] = []
        heading_count = 0
        
        para_index_for_detection = 0
        for paragraph in doc.paragraphs:
            para_index_for_detection += 1
            hit = detector.detect(paragraph, para_index=para_index_for_detection)
            paragraph_hits.append((paragraph, hit))
            if hit.is_heading:
                heading_count += 1
        
        # Определяем, нужен ли visual fallback
        # Если заголовков мало (< 3 на документ > 50 параграфов) → rerun с visual fallback
        total_paragraphs = len([p for p in doc.paragraphs if normalize_text(p.text)])
        enable_visual = False
        
        if heading_count == 0 or (total_paragraphs > 50 and heading_count < 3):
            enable_visual = True
            detector = HeadingDetector(enable_visual_fallback=True)
            detector.set_doc_stats(doc_stats)
            
            # Счётчик отклонений для numbering детекции (переиспользуем)
            detector._rejection_counter = rejection_counter
            
            # Пересчитываем hits с visual fallback
            paragraph_hits = []
            heading_count = 0
            para_index_for_detection = 0
            for paragraph in doc.paragraphs:
                para_index_for_detection += 1
                hit = detector.detect(paragraph, para_index=para_index_for_detection)
                paragraph_hits.append((paragraph, hit))
                if hit.is_heading:
                    heading_count += 1
        
        # Стек заголовков для построения section_path
        # Каждый элемент: (level, normalized_title)
        heading_stack: list[tuple[int, str]] = []
        
        # Флаг: был ли найден первый реальный заголовок (style или outline, не numbering/visual)
        first_real_heading_found = False
        
        # Счётчики ordinal для каждого (section_path, content_type)
        # Ключ: (section_path, content_type)
        ordinal_counters: dict[tuple[str, AnchorContentType], int] = {}
        
        anchors: list[AnchorCreate] = []
        warnings: list[str] = []
        
        # Диагностика заголовков
        heading_detected_count = 0
        heading_levels_histogram: dict[str, int] = {}
        heading_detection_mode_counts: dict[str, int] = {}
        # Подсчитываем общее количество отклонённых numbering кандидатов
        false_heading_filtered_count = sum(rejection_counter.values())
        frontmatter_paragraphs_count = 0  # Счётчик параграфов в frontmatter
        
        para_index = 0  # Счётчик всех параграфов для location_json
        
        # Проходим по всем параграфам документа с hits
        for paragraph, hit in paragraph_hits:
            para_index += 1
            
            # Получаем сырой текст
            text_raw = paragraph.text
            
            # Пропускаем пустые параграфы
            text_norm = normalize_text(text_raw)
            if not text_norm:
                continue
            
            # Определяем стиль
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Определяем content_type и обновляем стек заголовков
            content_type: AnchorContentType
            current_section_path: str
            
            # Проверяем, является ли это реальным заголовком (style или outline)
            is_real_heading = hit.is_heading and hit.mode in ("style", "outline")
            
            # Если это первый реальный заголовок, отмечаем это
            if is_real_heading and not first_real_heading_found:
                first_real_heading_found = True
            
            # Подсчитываем параграфы в frontmatter (до первого реального заголовка)
            if not first_real_heading_found:
                frontmatter_paragraphs_count += 1
            
            if hit.is_heading:
                # Это заголовок
                level = hit.level or 1
                
                # Обновляем стек заголовков только для реальных заголовков
                # Для numbering/visual заголовков до первого реального заголовка не обновляем стек
                if is_real_heading:
                    # Удаляем все заголовки с уровнем >= текущего уровня
                    heading_stack = [h for h in heading_stack if h[0] < level]
                    # Добавляем текущий заголовок (используем normalized_title)
                    heading_stack.append((level, hit.normalized_title))
                
                # Определяем section_path
                if first_real_heading_found:
                    # После первого реального заголовка используем нормальный путь
                    path_parts = [h[1] for h in heading_stack]
                    current_section_path = normalize_section_path(path_parts)
                else:
                    # До первого реального заголовка используем FRONTMATTER
                    current_section_path = FRONTMATTER_SECTION
                
                # Создаём HDR anchor только для реальных заголовков
                # Для numbering/visual заголовков до первого реального заголовка не создаём HDR
                if is_real_heading:
                    content_type = AnchorContentType.HDR
                    
                    # Обновляем диагностику
                    heading_detected_count += 1
                    level_str = str(level)
                    heading_levels_histogram[level_str] = heading_levels_histogram.get(level_str, 0) + 1
                    heading_detection_mode_counts[hit.mode] = heading_detection_mode_counts.get(hit.mode, 0) + 1
                else:
                    # Numbering/visual заголовок до первого реального заголовка → обрабатываем как обычный параграф
                    if is_list_item(paragraph):
                        content_type = AnchorContentType.LI
                    else:
                        content_type = AnchorContentType.P
                    # Не обновляем диагностику заголовков для таких случаев
            else:
                # Обычный параграф или элемент списка
                if is_list_item(paragraph):
                    content_type = AnchorContentType.LI
                else:
                    content_type = AnchorContentType.P
                
                # Определяем section_path
                if first_real_heading_found:
                    # После первого реального заголовка используем нормальный путь
                    path_parts = [h[1] for h in heading_stack]
                    current_section_path = normalize_section_path(path_parts)
                else:
                    # До первого реального заголовка используем FRONTMATTER
                    current_section_path = FRONTMATTER_SECTION
            
            # Получаем ordinal для данного (section_path, content_type) (используется только для поля ordinal, не для anchor_id)
            key = (current_section_path, content_type)
            ordinal = ordinal_counters.get(key, 0) + 1
            ordinal_counters[key] = ordinal
            
            # Вычисляем hash
            text_hash = get_text_hash(text_norm)
            
            # Формируем anchor_id: {doc_version_id}:{content_type}:{para_index}:{text_hash}
            # para_index используется вместо ordinal и section_path для стабильности при переносах
            doc_version_id_str = str(doc_version_id)
            anchor_id = f"{doc_version_id_str}:{content_type.value}:{para_index}:{text_hash}"
            
            # Формируем location_json
            location_json = {
                "para_index": para_index,
                "style": style_name,
                "section_path": current_section_path,
            }
            
            # Классифицируем source_zone для текущего section_path
            # Передаём section_path как список заголовков и текущий заголовок для лучшей классификации
            heading_text = None
            if heading_stack:
                heading_text = heading_stack[-1][1]  # Текст последнего заголовка
            
            # Определяем язык документа для классификатора
            doc_language = None
            if document_language in (DocumentLanguage.RU, DocumentLanguage.EN):
                doc_language = document_language.value
            
            zone_result = self.zone_classifier.classify(
                section_path=current_section_path,
                heading_text=heading_text,
                language=doc_language
            )
            source_zone = zone_result.zone
            
            # Определяем язык для anchor (локальная детекция из текста)
            anchor_language = detect_text_language(text_norm)
            
            # Создаём anchor
            anchor = AnchorCreate(
                doc_version_id=doc_version_id,
                anchor_id=anchor_id,
                section_path=current_section_path,
                content_type=content_type,
                ordinal=ordinal,
                text_raw=text_raw,
                text_norm=text_norm,
                text_hash=text_hash,
                location_json=location_json,
                source_zone=source_zone,
                language=anchor_language,
            )
            
            anchors.append(anchor)

        # Извлекаем footnotes из документа
        # В python-docx footnotes доступны через doc.part.footnotes_part.footnotes
        fn_section_path = "FOOTNOTES"
        footnotes_count = 0
        footnotes_anchors_count = 0
        
        try:
            # Получаем footnotes_part из документа
            footnotes_part = doc.part.footnotes_part if hasattr(doc.part, 'footnotes_part') else None
            
            if footnotes_part is not None:
                # Получаем коллекцию footnotes
                footnotes = footnotes_part.footnotes if hasattr(footnotes_part, 'footnotes') else None
                
                if footnotes is not None:
                    # Проходим по всем footnotes
                    for fn_idx, footnote in enumerate(footnotes):
                        footnotes_count += 1
                        
                        # Получаем параграфы внутри footnote
                        # В python-docx footnote имеет атрибут paragraphs
                        fn_paragraphs = footnote.paragraphs if hasattr(footnote, 'paragraphs') else []
                        
                        if not fn_paragraphs:
                            continue
                        
                        # Обрабатываем каждый параграф внутри footnote
                        for fn_para_idx, fn_paragraph in enumerate(fn_paragraphs, start=1):
                            # Получаем текст параграфа
                            text_raw = fn_paragraph.text if hasattr(fn_paragraph, 'text') else ""
                            
                            if not text_raw:
                                continue
                            
                            # Нормализуем текст
                            text_norm = normalize_text(text_raw)
                            if not text_norm:
                                continue
                            
                            # Получаем ordinal для поля ordinal (не используется в anchor_id)
                            key = (fn_section_path, AnchorContentType.FN)
                            ordinal = ordinal_counters.get(key, 0) + 1
                            ordinal_counters[key] = ordinal
                            
                            # Вычисляем hash
                            text_hash = get_text_hash(text_norm)
                            
                            # Формируем anchor_id для footnotes: {doc_version_id}:fn:{fn_index}:{fn_para_index}:{text_hash}
                            doc_version_id_str = str(doc_version_id)
                            anchor_id = f"{doc_version_id_str}:fn:{fn_idx}:{fn_para_idx}:{text_hash}"
                            
                            # Формируем location_json
                            location_json = {
                                "para_index": None,  # Для footnotes para_index не используется
                                "fn_index": fn_idx,  # Индекс сноски в документе
                                "fn_para_index": fn_para_idx,  # Индекс параграфа внутри сноски
                                "section_path": fn_section_path,
                            }
                            
                            # Классифицируем source_zone для footnotes (обычно unknown)
                            doc_language = None
                            if document_language in (DocumentLanguage.RU, DocumentLanguage.EN):
                                doc_language = document_language.value
                            
                            zone_result = self.zone_classifier.classify(
                                section_path=fn_section_path,
                                heading_text=None,
                                language=doc_language
                            )
                            
                            # Определяем язык для footnote anchor (локальная детекция из текста)
                            fn_anchor_language = detect_text_language(text_norm)
                            
                            # Создаём anchor для footnote
                            anchors.append(
                                AnchorCreate(
                                    doc_version_id=doc_version_id,
                                    anchor_id=anchor_id,
                                    section_path=fn_section_path,
                                    content_type=AnchorContentType.FN,
                                    ordinal=ordinal,
                                    text_raw=text_raw,
                                    text_norm=text_norm,
                                    text_hash=text_hash,
                                    location_json=location_json,
                                    source_zone=zone_result.zone,
                                    language=fn_anchor_language,
                                )
                            )
                            footnotes_anchors_count += 1
                    
                else:
                    # footnotes_part есть, но footnotes недоступны
                    # Это нормально, если в документе нет footnotes
                    pass
            else:
                # footnotes_part отсутствует - в документе нет footnotes
                # Это нормально, не все документы содержат footnotes
                pass
                
        except AttributeError as e:
            # Если API python-docx не поддерживает footnotes или структура отличается
            warnings.append(f"Footnotes недоступны через python-docx API: {str(e)}")
        except Exception as e:
            # Другие ошибки при извлечении footnotes
            warnings.append(f"Ошибка при извлечении footnotes: {str(e)}")
        
        # Определяем качество заголовков
        heading_quality = "none"
        if heading_detected_count == 0:
            heading_quality = "none"
            warnings.append("No headings detected; section_path fallback to ROOT")
        elif heading_detected_count <= 2:
            heading_quality = "low"
        else:
            # Проверяем, не слишком ли много visual fallback
            visual_count = heading_detection_mode_counts.get("visual", 0)
            if visual_count > 0 and visual_count / heading_detected_count > 0.8:
                heading_quality = "low"
                warnings.append("Headings detected mostly via visual fallback; verify structure")
            else:
                heading_quality = "ok"
        
        # Формируем summary
        counts_by_type = {}
        for anchor in anchors:
            content_type_str = anchor.content_type.value
            counts_by_type[content_type_str] = counts_by_type.get(content_type_str, 0) + 1
        
        # Собираем уникальные section_path
        unique_sections = set(anchor.section_path for anchor in anchors)
        
        summary = {
            "anchors_count": len(anchors),
            "counts_by_type": counts_by_type,
            "num_sections": len(unique_sections),
            "sections": sorted(list(unique_sections)),
            # Диагностика заголовков
            "heading_detected_count": heading_detected_count,
            "heading_levels_histogram": heading_levels_histogram,
            "heading_detection_mode_counts": heading_detection_mode_counts,
            "heading_quality": heading_quality,
            # Новые метрики фильтрации
            "false_heading_filtered_count": false_heading_filtered_count,
            "frontmatter_paragraphs_count": frontmatter_paragraphs_count,
            # Информация о footnotes
            "footnotes_count": footnotes_count,
            "footnotes_anchors_count": footnotes_anchors_count,
            "warnings": warnings,
        }
        
        return DocxIngestResult(
            anchors=anchors,
            summary=summary,
            warnings=warnings,
        )

